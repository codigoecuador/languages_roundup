import newspaper as np
from roundup_db1 import Entry, Category, Keyword, Publication, Author, Section
from roundup_db1 import DataAccessLayer
from dateutil.parser import parse
import itertools as it
import BTCInput2 as btc
from sqlalchemy import func
from sqlalchemy.orm.exc import MultipleResultsFound
import functools
import pprint
import warnings

import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
#from roundup_db2 import Entry

class Roundup(object):
    '''NOTE: do not try to use the roundup class for anything except exporting roundups'''
    @classmethod
    def prep_export(cls, title, start_date, end_date,
                   filename, min_category, max_category):
        new_sections = get_sections()
        new_sections = [RoundupSection.from_normal_section(i) for i in new_sections]
        result= cls(title=title, start_date=start_date, end_date=end_date,
                   filename=filename, min_category=min_category, max_category=max_category,
                   sections=new_sections)
        result.make_roundup()
        return result
    
    def __init__(self, title, start_date, end_date,
             filename, min_category=1,
             max_category=21, sections=[]):#app.get_sections()):
        self.title = title
        self.start_date = start_date
        self.end_date = end_date
        self.filename = filename
        self.min_category = min_category
        self.max_category = max_category
        self.sections = [RoundupSection.from_normal_section(i) for i in sections]
        self.categories = self.get_category_articles()
        self.introduction=None
    
    def make_roundup(self):
        '''Takes the sections and gets and organizes the categories. These data structures are used ONLY
        when we export the final roundup to docx or html. There is NO other use for them in the program logic.'''
        for section in self.sections:
            for category in self.categories:
                if category.section_id == section.section_id:
                    section.categories.append(category)
    
    @staticmethod
    def add_hyperlink(paragraph, text, url):
        # This gets access to the document.xml.rels file and gets a new relation id value
        #print(paragraph)
        #print(text)
        #print(url)
        try:
            part = paragraph.part
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

            # Create the w:hyperlink tag and add needed values
            hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
            hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

            # Create a w:r element and a new w:rPr element
            new_run = docx.oxml.shared.OxmlElement('w:r')
            rPr = docx.oxml.shared.OxmlElement('w:rPr')

            # Join all the xml elements together add add the required text to the w:r element
            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)

        # Create a new Run object and add the hyperlink into it
            r = paragraph.add_run ()
            r._r.append (hyperlink)

        # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
        # Delete this if using a template that has the hyperlink style in it
            r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
            r.font.underline = True

            return hyperlink
        except Exception as e:
            print(e)
            
    @staticmethod
    def add_article(document, article):
    #print(article)
        try:
            new_paragraph = document.add_paragraph('') #add blank paragraph that we append the text to
            Roundup.add_hyperlink(paragraph=new_paragraph, text=article.entry_name, url=article.entry_url)
            #print(Article.get_date_formatted(article))
            #new_paragraph.add_run(' ({0}) '.format(Entry.get_date_formatted(article))) 
            new_paragraph.add_run(f' ({article.get_date_formatted}) ')#.format(article.get_date_formatted))
            #blank space between the link and the description
            new_paragraph.add_run(article.description)
        except Exception as e:
            print(e)
    
    @staticmethod       
    def add_section(document, section):
        section_name = document.add_paragraph(section.name)
        section.categories.sort(key=lambda x: x.name, reverse=True)
        section.categories.reverse()
        for category in section.categories:
            Roundup.add_category(document, category)
    
    @staticmethod
    def add_category(document, category):
        category_name = document.add_paragraph(category.name)
        category.entries.sort(key=lambda x: x.entry_name, reverse=True)
        category.entries.reverse()
        for article in category.entries:
            Roundup.add_article(document, article)
    
    def create_roundup(self, document, roundup_title, sections):
        title = document.add_paragraph(roundup_title)
        for section in self.sections:
            Roundup.add_section(document, section)
            
    def get_category_articles(self):
        current_category = 1
        max_category = 21
        article_categories = []
        while current_category <= max_category:
    #for category in range(min_category, max_category+1):
            cat = get(session=dal.session, model=Category, category_id=current_category)
        #new_cat = RoundupCategory()
            print(cat)
            article_categories.append(RoundupCategory(category_id=cat.category_id,
                                                  name=cat.name,
                                                  entries = [i for i in cat.entries if i.date >=self.start_date and i.date <=self.end_date],
                                                     section_id=cat.section_id))
            current_category += 1
        return article_categories
        
    def export_docx(self):
        #def complete_roundup2(filename, roundup_title, sections):
        new_document = docx.Document()
        Roundup.create_roundup(self=self, document=new_document, roundup_title=self.title, sections=self.sections)
        new_document.save(f'{self.filename}.docx')#.format(self.filename))
        
class RoundupSection(object):
    '''Replicates the section class when we create roundups. DO NOT use this class
    for anything except roundup generation.'''
    @classmethod
    def from_normal_section(cls, section: Section):
        return cls(section_id = section.section_id,
                   name=section.name,
                   categories=[])
    
    def __init__(self, section_id, name, categories=[]):
        self.section_id = section_id
        self.name = name
        self.categories=categories
        #self.start_date = start_date
        #self.end_date = end_date

    def __repr__(self):
        return f'Sec(section_id={self.section_id}, name={self.name})'#.format(self.section_id, self.name)
                   
class RoundupCategory(object):
    '''Replicates the category class when we create roundups. Sqlalchemy methods
    cannot be called on this class. If we did not create a new class type like this,
    then we could not mess with the entries in that category without causing problems
    in the database as a whole. DO NOT use this class for anything except creating
    roundups.'''
    
    @classmethod
    def from_normal_category(cls, category: Category):
        return cls(category_id = category.category_id,
                  name=category.name, section_id = category.section_id, entries = [])
    
    def __init__(self, category_id, name, section_id, entries=[]):
        self.category_id = category_id
        self.name = name
        self.entries = entries
        #self.start_date = start_date
        #self.end_date = end_date
        self.section_id = section_id
        
    def __repr__(self):
        return f'ArtCat(category_id={self.category_id}, name={self.name})'#.format(self.category_id, self.name)

def export_roundup(title, start_date, end_date, filename, min_category=1, max_category=21):
    '''Exports a docx roundup'''
    new_roundup = Roundup.prep_export(title=title,
                    start_date=start_date, end_date=end_date, filename=filename, min_category=min_category, max_category=max_category)
    try:
        new_roundup.export_docx()
        print('Export successful')
    except Exception as e:
        print(e)

#def create_docx_roundup(line):
#    line=''
#    title = btc.read_text('Enter title or "." to return to main menu: ')
#    if title == '.': return
#    start_date = btc.read_date('Enter start date "(MM/DD/YYYY)": ')
#    end_date = btc.read_date('Enter end_date "MM/DD/YYYY": ')
#    filename = btc.read_text('Enter filename or "." to return to main menu: ')
    #call the export_roundup function to export a docx roundup
#    export_roundup(title=title, start_date=start_date, end_date=end_date,
                   #       filename=filename)

def create_docx_roundup(args):
    line=''
    if not args.title:
        title = btc.read_text('Enter title or "." to return to main menu: ')
        if title == '.': return
    else:
        title=' '.join(args.title)
    if not args.date_range:
        start_date = btc.read_date('Enter start date "(MM/DD/YYYY)": ')
        end_date = btc.read_date('Enter end_date "MM/DD/YYYY": ')
    else:
        start_date = parse(args.date_range[0]).date()
        end_date = parse(args.date_range[1]).date()
    if not args.filename:
        filename = btc.read_text('Enter filename or "." to return to main menu: ')
    #call the export_roundup function to export a docx roundup
    export_roundup(title=title, start_date=start_date, end_date=end_date,
                          filename=args.filename)
        
#create section

def add_category(session, category_name):
    result = get(session=session, model=Category, name=category_name)
    if result != None:
        print('Category exists')
        return
    else:
        new_category = Category.from_input(category_name=category_name, session=session)
        #new_category = Category(name=category_name)
        #print(new_category)
    confirm_choice = btc.read_int_ranged(f'Add {new_category.name}? 1 to add, 2 to cancel', 1, 2)
    if confirm_choice == 1:
        session.add(new_category)
        session.commit()
        print(f'{new_category.name} added to database')#.format(new_category.name))
    elif confirm_choice == 2:
        print('Category add cancelled')
        return

def add_item(session, search_type, new_name):
    """session is the current active session, search_type is the type of item, while
    new_name is the current type of item to add"""
    new_name=new_name.lower()
    search_types = {'author': Author, 'keyword': Keyword,
             'section': Section}
    result = get(session=session, model=search_types[search_type], name=new_name)
    if result != None:
        print('Item exists')
    else:
        confirm_choice = btc.read_int_ranged(f'Add {new_name}? 1 to add, 2 to cancel', 1, 2)
        if confirm_choice == 1:
            new_item=get_or_create(session=session, model=search_types[search_type], name=new_name)
            #session.add(new_keyword)
            #session.commit()
            print(f'{new_item} added to database')#.format(new_item))
        elif confirm_choice == 2:
            print('add cancelled')
            return
    
def add_pub_or_cat(session, add_type, new_name, second_item):
    '''use this one for the CLI'''
    add_type = add_type.lower()
    add_types = {'publication': Publication, 'category': Category}
    #item_twos = {'publication': 'url', 'category': 'section_id'}
    result = get(session=session, model=add_types.get(add_type), name=new_name, second_item=second_item)
    if result != None:
        print(f'{add_type} exists')#.format(add_type))
        print(result)
        return
    else:
        confirm_choice = btc.read_int_ranged(f'Add {new_name}? 1 to add, 2 to cancel', 1, 2)
        if confirm_choice == 1:
            get_or_create(session, model=add_types.get(add_type), name=new_name,
                      second_item=second_item)
        elif confirm_choice == 2:
            print(f'{add_type} add cancelled')#.format(add_type))
            return
        
def add_cat(session, line):
    '''Takes input from the CLI and passes it to the add_pub_or_cat function'''
    try:
        new_name, second_item = line.split()
        add_pub_or_cat(session=session, add_type='category',
                        new_name=new_name, second_item=second_item)
        return True
    except IndexError as e:
        print(e)
        return False

def add_pub(session, line):
    '''Takes input from the CLI and passes it to the add_pub_or_cat function'''
    try:
        new_name, second_item = line.split('|')
        add_pub_or_cat(session=session, add_type='publication',
                        new_name=new_name, second_item=second_item)
        return True
    except IndexError as e:
        print(e)
        return False
        
def add_publication(session, keyword_text):
    result = get(session=session, model=Keyword, word=keyword_text)
    if result != None:
        print('Publication exists')
        return
    else:
        new_keyword = Publication(pub_name=keyword_text)
    confirm_choice = btc.read_int_ranged(f'Add {new_keyword.pub_name}? 1 to add, 2 to cancel', 1, 2)
    if confirm_choice == 1:
        session.add(new_keyword)
        session.commit()
        print(f'{new_keyword.word} added to database')#.format(new_keyword.word))
    elif confirm_choice == 2:
        print('Keyword add cancelled')
        return
    
#modify articles
    
def add_keyword_to_article(session, entry_id):##model=Keyword):
    """Add a keyword to an existing article, the keyword is appended to the article's
    keywords attribute"""
    #new_keyword = btc.read_text('Enter new keyword: ')
    #make sure the keyword doesn't already exist
    entry_result = session.query(Entry).filter(Entry.entry_id==entry_id).scalar()
    #entry_result = entry_result.filter(Entry.entry_id==entry_id)
    #entry_result = entry_result.scalar()
    if entry_result != None:
        print('Entry found: ')
        print(entry_result)
        new_keyword=btc.read_text('Enter new keyword: ')
        edit_choice = btc.read_int_ranged('Add new keyword to this article? (1 for yes, 2 for no)', 1, 2)
        if edit_choice == 1:
            keyword_result = session.query(Keyword).filter(Keyword.word.like(f'%{new_keyword}%')).all()#.format(new_keyword))).all()
            if len(keyword_result) >= 1:
                print('Keyword exists')
                print(keyword_result)
                print('Entry found:')
                print(entry_result)
                keywords = it.chain(keyword_result)
                while True:
                        #we do this loop if the keyword exists
                    try:
                        item = next(keywords)
                        print(item)
                    except StopIteration:
                        print('No more keywords left')
                    item_choice = btc.read_int_ranged('Is this the keyword you want? (1-yes, 2-continue, 3-quit)', 
                                                      1, 3)
                    #1 select
                    if item_choice == 1:
                        try:
                            assert item not in entry_result.keywords
                        except AssertionError:
                            print('Keyword already attached to article')
                            print('Returning to main menu')
                            return
                        entry_result.keywords.append(item)
                        session.commit()
                        print('Keyword added successfully')
                        break
                    elif item_choice == 2:
                            #continue
                        continue
                    elif item_choice == 3:
                        print('Keyword add cancelled, return to main menu')
                        return
            elif len(keyword_result) ==0:
                print('Keyword does not exist')
                #new_keyword = btc.read_text('Enter new keyword: ')
                kw = Keyword(word=new_keyword)
                make_keyword_choice = btc.read_int_ranged(f'Create {kw} as a new keyword for ? {entry_result.entry_name} (1 yes, 2 no)',1, 2)#.format(kw,
                                                        #entry_result.entry_name), 1, 2)
                if make_keyword_choice == 1:
                    entry_result.keywords.append(kw)
                    session.commit()
                    print('Keyword add completed')
                elif make_keyword_choice == 2:
                    print('Add keyword cancelled')
                    return
        elif edit_choice == 2:
            print('Keyword edit cancelled, returning to main menu')
            return
    elif entry_result == None:
        print('Entry not found, returning to main menu')
        return

def delete_entry_keyword(session, entry_id):
    '''Delete a keyword from an existing article by popping it from the keyword list'''
    entry_result = session.query(Entry).filter(Entry.entry_id==entry_id).scalar()
    if entry_result != None:
        article_keywords = it.cycle(entry_result.keywords)
        while True:
            pprint.pprint(entry_result.keywords)
            activeItem = next(article_keywords)
            #print('Active item type:', type(activeItem))
            print('Enter 1 to delete keyword, 2 to continue, 3 to exit to main menu')
            delete_choice = btc.read_int_ranged(f'Delete {activeItem} from the keywords?', 1, 3)#.format(activeItem),1, 3)
            if delete_choice == 1:
                entry_result.keywords.remove(activeItem)
                session.commit()
                #print(entry_result.keywords)
            elif delete_choice == 2:
                continue
            elif delete_choice == 3:
                print('Returning to main menu')
                break
    else:
        print('Not found, return to main menu')
        return

    

def make_article(article: str) -> np.Article:
    article = np.Article(article)
    article.build()
    return article

def get_or_create(session, model, **kwargs):
    '''If an object is present in the database, it returns the object.
    Otherwise, it creates a new instance of that object'''
    instance = session.query(model).filter_by(**kwargs).first()
    if instance:
        return instance
    else:
        instance = model(**kwargs)
        session.add(instance)
        session.commit()
        return instance

#read section - functions associated with reading items]

    
def get(session, model, **kwargs):
    '''If an object is present in the database, it returns the object.
    Otherwise, it creates a new instance of that object'''
    instance = session.query(model).filter_by(**kwargs).scalar()
    return instance

def check_date(article):
    '''Check to see if the date is there in the article'''
    try:
        assert article.publish_date, 'No publish date found'
        return article.publish_date
    except AssertionError:
        return None
    
def create_date(article):
    '''Checks to see if the date is there by calling new_date,
    then if there is no date, it asks the user for input'''
    new_date = check_date(article)
    if new_date == None:
        dateObj = input('Enter the date mm/dd/yyyy')
        new_date = parse(dateObj)
        #return new_date
    else:
        new_date = input('Enter date (mm/dd/yyyy): ')
        new_date = parse(new_date)
    return new_date    
    
def display_categories(section_id=None):
    query = dal.session.query(Category)
    if section_id != None:
        query = query.filter(Category.section_id == section_id)
    query = query.all()
    print('Categories: ')
    cat_map = map(str, query)
    print('\n'.join(cat_map))
    #for i in query:
    #    print(i.category_id, i.name, end='  ')

def display_sections():
    query = dal.session.query(Section).all()
    section_map = map(str, query)
    print('Sections: ')
    print('\n'.join(section_map))

def get_description():
    description = input('Enter article description (max 500 characters): ')
    return description
        
def get_sections():
    '''Return a list of all the sections in the program'''
    query = dal.session.query(Section).all()
    return query

def get_categories():
    '''Return a list of all the categories in the program'''
    query = dal.session.query(Category).all()
    return query

def count_articles(line, session):
    '''Takes the input from article_count in cmd_roundup.ipynb and determines
    if there is more than one article.'''
    try:
        start_date, end_date = line.split() #tuple unpacking means the first and last elements are taken
        date_range_count(start_date=start_date, end_date=end_date,
                                session=session)
    except ValueError:
        article_count(session=session)

def article_count(session):
    query = session.query(Category.name, func.count(Entry.entry_id))
    query= query.outerjoin(Entry).group_by(Category.name)
    query = query.order_by(func.count(Entry.entry_id))
    query = query.all()
    for row in query[::-1]:
        print(row)
    undescribed_articles=session.query(Entry).filter(Entry.description.like('%not specified%')).all()
    print('Undescribed articles', len(undescribed_articles))
        
def date_range_count(start_date, end_date, session):
    '''Combine date_range_count and article_count'''
    #print(line)
    try:
        #line = line.split(' ')
        start_date, end_date = parse(start_date), parse(end_date)
        start_date = start_date.date()
        end_date = end_date.date()
        print('start date:', start_date)
        print('end date:', end_date)
    except IndexError as e:
            print(e)
            return
    query = session.query(Category.name, func.count(Entry.entry_id))
    query= query.outerjoin(Entry).group_by(Category.name)
    query = query.filter(Entry.date >= start_date, end_date <= end_date)
    #query = query.filter(Entry.date <= end_date)
    query = query.order_by(func.count(Entry.entry_id))
    query = query.all()
    undesc = session.query(Entry).filter(Entry.date >= start_date) #get number of undescribed articles
    undesc = undesc.filter(Entry.date <= end_date)
    undesc = undesc.filter(Entry.description.like('%not specified%')).all()
    undesc = len(undesc)
    total = functools.reduce(lambda x, y: x+y,[row[1] for row in query])
    for row in query[::-1]:
        print(row)
    print(total, f'articles total from {start_date} to {end_date}')
    print(f'Undescribed articles: {undesc}')
    
def articles_needed(start_date, end_date, session):
    min_articles_cat = 5
    '''Combine date_range_count and article_count'''
    #print(line)
    try:
        #line = line.split(' ')
        start_date, end_date = parse(start_date), parse(end_date)
        start_date = start_date.date()
        end_date = end_date.date()
        print('start date:', start_date)
        print('end date:', end_date)
    except IndexError as e:
            print(e)
            return
    query = session.query(Category.name, func.count(Entry.entry_id))
    query= query.outerjoin(Entry).group_by(Category.name)
    query = query.filter(Entry.date >= start_date, end_date <= end_date)
    #query = query.filter(Entry.date <= end_date)
    query = query.order_by(func.count(Entry.entry_id))
    query = query.all()
    #total = functools.reduce(lambda x, y: x+y,[row[1] for row in query])
    articles_needed = {}
    for row in query[::-1]:
        #print(row)
        if row[1] < 5:
            #print(row, '')
            articles_needed[row[0]] = (min_articles_cat - row[1])
    pprint.pprint(articles_needed)
    #print(articles_needed)
    #for k, v in articles_needed:
        #print('{0}: {1} articles, {2} more articles needed'.format(k, v[0], v[1]))
    #print(total, 'articles total from {0} to {1}'.format(start_date, end_date))
    
def find_entry(args, session):
    if args.entry_id and args.id_range:
        raise Exception('Must be either id or id range')
    elif args.date and args.date_range:
        raise Exception('Must be either date or date range')
    else:
        query = session.query(Entry)
        if args.category_id:
            query = query.filter(Entry.category_id == args.category_id)
        if args.id_range:
            query = query.filter(Entry.entry_id >= args.id_range[0],
                                    Entry.entry_id <= args.id_range[1])
        if args.entry_id:
            query = query.filter(Entry.entry_id == args.entry_id)
        if args.date:
            date = parse(args.date).date()
            query = query.filter(Entry.date == date)
        if args.date_range:
            query = query.filter(Entry.date >= parse(args.date_range[0]).date(),
                                    Entry.date <= parse(args.date_range[1]).date())
        if args.url:
            query = query.filter(Entry.url.like(f'%{args.url}%'))
        if args.title:
            query = query.filter(Entry.entry_name.like(f'%{args.title}%'))
        result = query.all()
        result_total = len(result)
        if result_total == 0:
            print('no entries found')
            return
        result_cycle = it.cycle(result)
        print(f'{result_total} entries found')
        info_choice = btc.read_int_ranged('1 to view results, 2 to cancel: ', 1, 2)
        if info_choice == 1:
            while True:
                continue_choice = btc.read_int_ranged('1 to view next, 2 to quit', 1, 2)
                print(next(result_cycle).name)
                if continue_choice == 1:
                    print(next(result_cycle))
                elif continue_choice == 2:
                    print('returning to main menu')
                    break

def find_category(args, session):
    query = session.query(Category)
    if args.category_id:
        query = query.filter(Category.category_id == args.category_id)
    if args.category_name:
        print(args.category_name)
        category_name = ' '.join(args.category_name)
        print(category_name)
        query = query.filter(Category.name.like(f"%{category_name}%"))
    if args.section_id:
        query = query.filter(Category.section_id == args.section_id)
    result = query.all()
    result_total = len(result)
    if result_total == 0:
        print('no categories found')
        return
    result_cycle = it.cycle(result)
    print(f'{result_total} categories found')
    info_choice = btc.read_int_ranged('1 to view results, 2 to cancel: ', 1, 2)
    if info_choice == 1:
        while True:
            continue_choice = btc.read_int_ranged('1 to view next, 2 to quit: ', 1, 2)
            print(next(result_cycle).name)
            if continue_choice == 1:
                print(next(result_cycle))
            elif continue_choice == 2:
                print('returning to main menu')
                break                    
               
                    
def search_exact_name(line, session):
    search_types = {'entry': Entry, 'category': Category, 'publication': Publication,
                   'section': Section, 'keyword':Keyword}
    #app.search_exact_name(line=line)
    line = line.split(' ')
    search_type = line[0].lower() #make it lowercase to fit in the dictionary
    value= ' '.join(line[1:]) #The first word is the search type, so we want the remainder of the words
        #to be the title
    if search_type in search_types:
        #print(True)
        result = get(session=session, model=search_types[search_type], name_value=value)
        print(result)
    else:
        print('Invalid search type, return to main menu')

def search_by_id(search_type, item_id, session):
    '''This will serve as a universal function to get an item by its id'''
    search_type = search_type.lower()
    search_types = {'entry': Entry, 'category': Category,
                   'publication': Publication, 'section': Section,
                   'keyword': Keyword, 'author': Author}
    item_types = {'entry': 'keywords and authors', 'category': 'entries',
                   'publication': 'entries', 'section': 'categories',
                   'keyword': 'entries', 'author': 'entries'}
    if search_type in search_types:
        result = get(session=session, model=search_types[search_type], id_value=item_id)
        print(result)
        info_choice = btc.read_int_ranged('View more information? (1-yes, 2-quit) ',1,2)
        if info_choice == 1:
            misc = it.cycle(result.items)
            while True:
                print(f'Cycles through {item_types[search_type]}')#.format(item_types[search_type]))
                continue_choice = btc.read_int_ranged('1 to view next, 2 to quit', 1, 2)
                if continue_choice == 1:
                    print(next(misc))
                elif continue_choice == 2:
                    print('returning to main menu')
                    break
    else:
        print('Invalid search type. Return to main menu.')

def name_search(session, line):
    category_result = session.query(Category).filter(Category.name.like(f'%{line}%'))#.format(line))).all()
    section_result = session.query(Section).filter(Section.name.like(f'%{line}%'))#.format(line))).all()
    entry_result = session.query(Entry).filter(Entry.entry_name.like(f'%{line}%'))#.format(line))).all()
    keyword_result = session.query(Keyword).filter(Keyword.word.like(f'%{line}%')).all()
    publication_result = session.query(Publication).filter(Publication.title.like(f'%{line}%')).all()
    author_result = session.query(Author).filter(Author.author_name.like(f'%{line}%')).all()
    result = it.cycle(category_result+section_result+entry_result+keyword_result+publication_result+author_result)
    #result2 = it.cycle(result)
    while True:
        try:
            pprint.pprint(next(result))
        except StopIteration:
            print('no more left')
        continue_choice = btc.read_int_ranged('1 to view more results, 2 to return to main menu', 1,2)
        if continue_choice != 1:
            break
            
def name_search2(session, line):
    category_result = session.query(Category).filter(Category.name.like(f'%{line}%')).all()
    section_result = session.query(Section).filter(Section.name.like(f'%{line}%')).all()#.format(line))).all()
    entry_result = session.query(Entry).filter(Entry.entry_name.like(f'%{line}%')).all()#.format(line))).all()
    keyword_result = session.query(Keyword).filter(Keyword.word.like(f'%{line}%')).all()#.format(line))).all()
    publication_result = session.query(Publication).filter(Publication.title.like(f'%{line}%')).all()#like('%{0}%'.format(line))).all()
    author_result = session.query(Author).filter(Author.author_name.like(f'%{line}%')).all()
    result = it.cycle(category_result+section_result+entry_result+keyword_result+publication_result+author_result)
    #result2 = it.cycle(result)
    while True:
        try:
            pprint.pprint(next(result))
        except StopIteration:
            print('no more left')
        continue_choice = btc.read_int_ranged('1 to view more results, 2 to return to main menu', 1,2)
        if continue_choice != 1:
            break

def get_articles_for_roundup(start_date, end_date, category_id):
    '''
    Do not mess with this function without absolute certainty that you will
    not break the roundup generation process.
    '''
    query = dal.session.query(Category)
    query = query.filter(Category.category_id == category_id)
    query = query.first()
    return query

def get_entries_by_category(session, line):
    result = session.query(Category)
    result = result.filter(Category.name.like(f'%{line}%'))
    try:
        result = result.one()
    except MultipleResultsFound:
        print('Multiple results found')
        potentials = it.cycle(result)
        while True:
            potential_result = next(potentials)
            print(potential_result)
            result_choice = btc.read_bool(decision='is this the category? (y-yes, n-no): ',
                                         yes='y', no='n',
                                         yes_option='select', no_option='continue')
            if result_choice == True:
                result = potential_result
                break
            else:
                continue
                
    print(result)
    review_choice = btc.read_bool(decision=f'View articles from {result}',
                                       yes='y', no='n',
                                       yes_option='continue', no_option='cancel')
    print(review_choice)
    if review_choice == True:
        entries_by_cat = it.cycle(result.entries)
        while True:
            continue_choice = btc.read_bool(decision=f'View next article from {result}',
                                    yes='y', no='n',
                                    yes_option='continue', no_option='cancel')
            if continue_choice == True:
                print(next(entries_by_cat))
            elif continue_choice == False:
                print('Returning to main menu')
                break

#Generic EDIT section: these functions relate to editing any item

def edit_name(session, model, id_value, new_name):
    '''Works for: Entry, Category, Section, Keyword, Publication, Author'''
    models = {'entry': Entry, 'category': Category, 'author': Author,
             'section': Section, 'publication': Publication}
    if model in models:
        query = session.query(models.get(model,
            'Please enter "entry", category, section, keyword, publication, author'))
        query = query.filter(models.get(model).id_value == id_value)
        result = query.one()
        result.name_value = new_name
        session.commit()
    else:
        return
    
def edit_second_item(session, model, id_value, new_second_value):
    """Works for Publication, Category"""
    models = {'publication': Publication, 'category': Category}
    if model in models:
        query = session.query(models.get(model,
            'Please enter "category" or "publication"'))
        query = query.filter(models.get(model).id_value == id_value)
        result = query.one()
        result.second_item = new_second_value
        session.commit()
    else:
        return
        
        

#Edit category section: these functions relate to editing categories

#Edit keyword section: formatting keywords

#Edit publication section: editing publications

#functions that capture input for a publication

#Edit section section: edit sections

#finalize section: functions to finalize articles

def finalize(session, start_date, end_date):
    start_date = parse(start_date)
    end_date = parse(end_date)
    query = session.query(Entry).filter(Entry.date >= start_date, Entry.date <= end_date)
    query = query.filter(Entry.description.like('%not specified%')).all()
    #print(query)
    result = it.cycle(query)
    undescribed_articles = len(query)
    print(f'{undescribed_articles} undescribed articles')
    while True:
        try:
            active_item=next(result)
        except StopIteration:
            print('No undescribed entries, return to main menu')
            return
        print('Next entry: ', active_item.name)
        continue_choice = btc.read_int_ranged('press 1 to continue, 2 to quit', 1, 2)
        if continue_choice == 1:
            #os.system('cls||clear')
            undesc = session.query(Entry).filter(Entry.date >= start_date, Entry.date <= end_date)
            undesc = undesc.filter(Entry.description.like('%not specified%')).all()
            undescribed = len(undesc)
            print(f'{undescribed} undescribed articles remaining')
            #active_item = next(result)
            print(f'''\n
Entry ID: {active_item.id_value}
Title: {active_item.name_value}
Date: {active_item.date}
Link: {active_item.entry_url}
Authors: {active_item.authors}
Publication: {active_item.publication}
Category: {active_item.category}
Description: {active_item.description}
Keywords: {active_item.keywords}''')
            #print(active_item, end='\n')
            edit_choice = btc.read_int_ranged('Edit description 1-yes, 2-no, 3-quit? ', 1, 3)
            if edit_choice == 1:
                #new_desc = get_description()
                new_desc = btc.read_text('Enter new description or "." to cancel: ')
                if new_desc != '.': 
                    active_item.description = new_desc
                    session.commit()
                else:
                    new_desc = 'Not specified' #if the user doens't edit it
                #undescribed_articles -= 1
            elif edit_choice ==2:
                continue
            elif edit_choice == 3:
                print('Returning to main menu')
                return
        elif continue_choice == 2:
            break
            
def finalize2(session, start_date, end_date):
    start_date = parse(start_date)
    end_date = parse(end_date)
    query = session.query(Entry).filter(Entry.date >= start_date, Entry.date <= end_date)
    query = query.filter(Entry.description.like('%not specified%')).all()
    #print(query)
    result = it.cycle(query)
    undescribed_articles = len(query)
    print(f'{undescribed_articles} undescribed articles')
    while True:
        try:
            active_item=next(result)
        except StopIteration:
            print('No undescribed entries, return to main menu')
            return
        print('Next entry: ', active_item.name)
        continue_choice = btc.read_int_ranged('press 1 to continue, 2 to quit', 1, 2)
        if continue_choice == 1:
            #os.system('cls||clear')
            undesc = session.query(Entry).filter(Entry.date >= start_date, Entry.date <= end_date)
            undesc = undesc.filter(Entry.description.like('%not specified%')).all()
            undescribed = len(undesc)
            print(f'{undescribed} undescribed articles remaining')
            #active_item = next(result)
            while True:
                print(f'''\n
Entry ID: {active_item.id_value}
Title: {active_item.name_value}
Date: {active_item.date}
Link: {active_item.entry_url}
Authors: {active_item.authors}
Publication: {active_item.publication}
Category: {active_item.category}
Description: {active_item.description}
Keywords: {active_item.keywords}''')
            #print(active_item, end='\n')
                edit_menu = ('''
1. Edit description
2. Edit category id''')
                edit_choice = btc.read_int_ranged('Edit description - 1, Edit category id - 2, 3-next_article: ', 1, 3)
                if edit_choice == 1:
                    #new_desc = get_description()
                    #desc_finalize(entry_id=active_item.id_value, session=session)
                    summary_choice = btc.read_int_ranged('Type 1 to view summary, 2 to skip', 1, 2)
                    if summary_choice == 1:
                        print(f'Summary:\n{active_item.summary}')
                    else:
                        print('Summary display not needed')
                    new_desc = btc.read_text('Enter new description or "." to cancel: ')
                    print(new_desc)
                    if new_desc != '.': 
                        active_item.description = new_desc
                        session.commit()
                    else:
                        new_desc = 'Not specified' #if the user doens't edit it
                elif edit_choice == 2:
                    cat_id_finalize(entry_id=active_item.id_value, session=session)
                elif edit_choice == 3:
                    break
        elif continue_choice == 2:
            break

#Edit entry section: these functions relate to editing entries

def name_from_input(session, entry_id):
    '''Gets user input for the article name, the name editing itself is carried out by the edit_name function'''
    query = session.query(Entry)
    query = query.filter(Entry.entry_id == entry_id)
    result = query.one()
    print(result)
    edit_choice = btc.read_int_ranged('Edit title (1 for yes, 2 for no): ', 1, 2)
    if edit_choice == 1:
        new_name = btc.read_text('Enter new title or "." to cancel: ')
        if new_name != '.':
            #edit_name(session=session, entry_id=entry_id, new_name=new_name)
            edit_name(session=session, model='entry', id_value = entry_id, new_name=new_name)
        else:
            print('Edit description cancelled')
            return
    if edit_choice == 2:
        print('Edit cancelled')
        return
    
def edit_date(session, entry_id, new_date):
    query = session.query(Entry)
    query = query.filter(Entry.entry_id == entry_id)
    result = query.one()
    result.date = new_date
    session.commit()
    print('Entry edit successful')
    
def date_from_input(session, entry_id):
    query = session.query(Entry)
    query = query.filter(Entry.entry_id == entry_id)
    result = query.one()
    print(result)
    edit_choice = btc.read_int_ranged('Edit date (1 for yes, 2 for no): ', 1, 2)
    if edit_choice == 1:
        new_date = btc.read_text('Enter new date or "." to cancel: ')
        new_date = parse(new_date)
        if new_date != '.':
            edit_date(session=session, entry_id=entry_id, new_date=new_date)
        else:
            print('Edit date cancelled')
            return
    if edit_choice == 2:
        print('Edit cancelled')
        return
    
def edit_description(session, entry_id, new_description):
    query = session.query(Entry)
    query = query.filter(Entry.entry_id == entry_id)
    result = query.one()
    result.description = new_description
    session.commit()
    print('Entry edit successful')
    
def desc_from_input(session, entry_id):
    query = session.query(Entry)
    query = query.filter(Entry.entry_id == entry_id)
    result = query.one()
    print(result)
    edit_choice = btc.read_int_ranged('Edit description (1 for yes, 2 for no): ', 1, 2)
    if edit_choice == 1:
        new_description = btc.read_text('Enter new description or "." to cancel: ')
        if new_description != '.':
            edit_description(session=session, entry_id=entry_id, new_description=new_description)
        else:
            print('Edit description cancelled')
            return
    if edit_choice == 2:
        print('Edit cancelled')
        return
    
def desc_finalize(session, entry_id):
    edit_choice = btc.read_int_ranged('Edit description (1 for yes, 2 for no): ', 1, 2)
    if edit_choice == 1:
        new_description = btc.read_text('Enter new description or "." to cancel: ')
        if new_description != '.':
            edit_description(session=session, entry_id=entry_id, new_description=new_description)
        else:
            print('Edit description cancelled')
            return
    if edit_choice == 2:
        print('Edit cancelled')
        return

def cat_id_finalize(session, entry_id):
    display_categories()
    new_category_id = btc.read_int('Enter new category id or 0 to cancel: ')
    if new_category_id == 0:
        print('edit category id cancelled')
    else:
        query = session.query(Entry)
        query = query.filter(Entry.entry_id == entry_id)
        result = query.one()
        result.category_id = new_category_id
        session.commit()
        print('Entry edit successful')    
    
def edit_category_id(session, entry_id, new_category_id):
    query = session.query(Entry)
    query = query.filter(Entry.entry_id == entry_id)
    result = query.one()
    result.category_id = new_category_id
    session.commit()
    print('Entry edit successful')
    
def cat_id_from_input(session, entry_id):
    query = session.query(Entry)
    query = query.filter(Entry.entry_id==entry_id)
    result = query.one()
    print(result)
    edit_choice=btc.read_int_ranged('Edit category ID (1 for yes, 2 for no): ', 1, 2)
    if edit_choice == 1:
        display_categories()
        new_category_id = btc.read_int('Enter category id or 0 to cancel: ')
        if new_category_id <= 0:
            print('Edit category ID cancelled, return to main menu')
            return
        else:
            edit_category_id(session=session, entry_id=entry_id,
                             new_category_id=new_category_id)
            
    
def edit_entry(session, entry_id):
    entry = get(session=session, model=Entry, entry_id=entry_id)
    options = it.cycle([name_from_input, date_from_input,
                        desc_from_input, cat_id_from_input])
    while True:
        choice = btc.read_int_ranged(f'Continue editing {entry.entry_name}, 1-yes 2-quit', 1, 2)
        if choice == 1:
            next(options)(session, entry_id)
        elif choice == 2:
            break
            
#create section - functions for creating new instances

#add_articles: this section is for adding entries, this is separate due to the complexity of the code. The entries
#must be downloaded using the newspaper app.

def add_entry(session, url, category_id=None, date=None):
    new_entry = from_newspaper_two(url=url, category_id=category_id, date=date)
    session.add(new_entry)
    session.commit()
    print(f'{new_entry.entry_name} added to database')#.format(new_entry.entry_name))

def qa(session, url, category_id, date, description=None):
    'qa is short for quick_add'
    try:
        assert len(category_id) <=4, 'category_id out of order'
        assert parse(date), 'date out of order'
        assert len(url) > 10, 'url out of order'
    
    except AssertionError as e:
        #one of the parameters is entered incorrectly
        print(e)
        return
    new_article = make_article(url)
    print('\nTitle is being added...')#, new_article.title)
    pprint.pprint(new_article.title)
    new_pub = get(session, Publication,
                               url=new_article.source_url)
    print(new_pub)
    if new_pub == None:
        new_pub = get_or_create(session, Publication,
                        title=new_article.source_url,
                        url=new_article.source_url)
        print(new_pub)
    date = parse(date)
    if description == None:
        description='Not specified'
    authors = [get_or_create(session, Author, author_name=i) for i in new_article.authors]
    keywords = [get_or_create(session, Keyword, word=i) for i in new_article.keywords]
    new_entry= create_entry(article=new_article, description=description,
                            publication_id=new_pub.publication_id, category_id=category_id,
                            date=date, authors=authors, keywords=keywords)
    session.add(new_entry)
    session.commit()
    print(f'{new_entry.name} added successfully')#.format(new_entry.name))
    
def from_newspaper(url):
    #include confirm option to make sure that the user wants to add the article
    new_article = make_article(url)
    print('\nTitle:')#, new_article.title)
    pprint.pprint(new_article.title)
    print('Summary:')#, new_article.summary)
    pprint.pprint(new_article.summary)
    try:
        new_pub = get(dal.session, Publication,
                               url=new_article.source_url)
        if new_pub != None:
            print(new_pub)
            pub_choice = btc.read_bool(decision=f'Confirm {new_pub} as article publication? ',
                                       yes='y', no='n',
                                       yes_option='confirm', no_option='cancel')
            if pub_choice != True:
                #if the user REJECTS the publication that is listed
                new_title = input('Enter publication title: ')
                new_source_url = input('Enter source URL: ')
                new_pub = get_or_create(dal.session, Publication,
                                title=new_title,
                               url=new_source_url)
        else:
            #if we have a source URL but no title
            new_title = input('Enter publication title: ')
            new_pub = get_or_create(dal.session, Publication,
                                title=new_title,
                               url=new_article.source_url)
    except Exception as e:
        #if there's already a publication with that URL
        new_title = input('Enter publication title: ')
        new_source_url = input('Enter source URL: ')
        new_pub = get_or_create(dal.session, Publication,
                                title=new_title,
                               url=new_source_url)
    publication_id = new_pub.publication_id
    print('-'*64)
    display_categories()
    category_id = int(input('Enter category ID: '))
    #pub_title = input('Enter publication title: ')
    
    date = create_date(new_article)
    description=get_description()
    confirm_choice = btc.read_int_ranged('Confirm article add (1-yes, 2-no) : ', 1, 2)
    if confirm_choice == 2:
        print('Article add cancelled')
        return
    else:
        authors = [get_or_create(dal.session, Author, author_name=i) for i in new_article.authors]
        keywords = [get_or_create(dal.session, Keyword, word=i) for i in new_article.keywords]
        return create_entry(article=new_article, description=description,
                            publication_id=publication_id, category_id=category_id,
                            date=date, authors=authors, keywords=keywords)
    
def from_newspaper_two(url, category_id = None, date=None):
    #include confirm option to make sure that the user wants to add the article
    new_article = make_article(url)
    print('\nTitle:')#, new_article.title)
    pprint.pprint(new_article.title)
    print('Summary:')#, new_article.summary)
    pprint.pprint(new_article.summary)
    try:
        new_pub = get(dal.session, Publication,
                               url=new_article.source_url)
        if new_pub != None:
            print(new_pub)
            pub_choice = btc.read_bool(decision=f'Confirm {new_pub} as article publication? ',
                                       yes='y', no='n',
                                       yes_option='confirm', no_option='cancel')
            if pub_choice != True:
                #if the user REJECTS the publication that is listed
                new_title = input('Enter publication title: ')
                new_source_url = input('Enter source URL: ')
                new_pub = get_or_create(dal.session, Publication,
                                title=new_title,
                               url=new_source_url)
        else:
            #if we have a source URL but no title
            new_title = input('Enter publication title: ')
            new_pub = get_or_create(dal.session, Publication,
                                title=new_title,
                               url=new_article.source_url)
    except Exception as e:
        #if there's already a publication with that URL
        new_title = input('Enter publication title: ')
        new_source_url = input('Enter source URL: ')
        new_pub = get_or_create(dal.session, Publication,
                                title=new_title,
                               url=new_source_url)
    publication_id = new_pub.publication_id
    print('-'*64)
    if category_id == None:
        display_categories()
        category_id = int(input('Enter category ID: '))
    #pub_title = input('Enter publication title: ')
    if date == None:
        date = create_date(new_article)
    description=get_description()
    confirm_choice = btc.read_int_ranged('Confirm article add (1-yes, 2-no) : ', 1, 2)
    if confirm_choice == 2:
        print('Article add cancelled')
        return
    else:
        authors = [get_or_create(dal.session, Author, author_name=i) for i in new_article.authors]
        keywords = [get_or_create(dal.session, Keyword, word=i) for i in new_article.keywords]
        return create_entry(article=new_article, description=description,
                            publication_id=publication_id, category_id=category_id,
                            date=date, authors=authors, keywords=keywords)



def create_entry(article, description, publication_id, category_id, date, authors=[],
                 keywords=[]):
    #date=parse(date)
    return Entry(entry_name=article.title, entry_url=article.url,
          description=description, summary=article.summary,
        authors=authors,
          publication_id=publication_id,
    category_id=category_id,
        keywords=keywords,
            date=date)

#delete section - delete article

def delete_item(session, model, id_value):
    model = model.lower()
    models = {'entry': Entry, 'category': Category, 'keyword': Keyword,
             'author': Author, 'publication': Publication, 'section': Section}
    result = get(session=session, model = models.get(model, 'invalid delete type'),
                id_value=id_value)
    if result != None:
        print(result)
        delete_choice = btc.read_int_ranged(f'Delete {model} (1 for yes, 2 for no)?',
                                            1, 2)
        if delete_choice == 1:
            if (model == 'category') or (model=='section') or (model=='publication'):
                try:
                    assert len(result.items) == 0
                except AssertionError:
                    print('result has items, delete these first: ', result.items)
                    return
            confirm_choice = btc.read_int_ranged('Are you sure (1 for yes, 2 for no)?', 1, 2)
            if confirm_choice == 1:
                #delete the article
                session.delete(result)
                session.commit()
                print(f'{model} deleted')#.format(model))
            elif confirm_choice == 2:
                print('Delete cancelled')
                print(f'{result.name_value} remains in database')#.format(result.name_value))
        else:
            print('Delete cancelled by user, returning to main menu')
    else:
        print('Item not found, delete cancelled')
    
def export_html(session, program, start_date, end_date, title):
    filename = program + '.html'
    f = open(filename, 'w')

    opening_wrapper = f"""<html>
    <head>
    <title>{title}</title>
    </head>
    <body><p>{title}</p>"""
    f.write(opening_wrapper)
    section_query = session.query(Section)
    section_query = section_query.all()
    for section in section_query:
        f.write(section.wrapped_html_string)
        for category in section.categories:
            f.write(category.wrapped_html_string)
            for entry in category.entries:
                if (entry.date >= start_date) and (entry.date <= end_date):
                    f.write(entry.wrapped_html_string)
    closing_wrapper = """</body>
    </html>"""
    f.write(closing_wrapper)

def make_html_roundup(line, session):
    del line
    filename=btc.read_text('Please enter filename or "." to cancel: ')
    if filename == ".": return
    title=btc.read_text('Please enter title or "." to cancel: ')
    if title == ".": return
    start_date = btc.read_date('Pease enter start date ("MM/DD/YYYY"): ')
    end_date = btc.read_date('Please enter end date ("MM/DD/YYYY"):')
    try:
        export_html2(session=session, program=filename, title=title,
                           start_date=start_date, end_date=end_date)
        print(f'{filename} exported successfully')#.format(filename) )
    except Exception as e:
        print(e)
    
def export_html2(session, program, start_date, end_date, title):
    filename = program + '.html'
    f = open(filename, 'w')

    opening_wrapper = f"""<html>
    <head>
    <title>{title}</title>
    </head>
    <body><p>{title}</p>"""#.format(title)
    f.write(opening_wrapper)
    section_query = session.query(Section)
    section_query = section_query.all()
    for section in section_query:
        f.write(section.wrapped_html_string)
        for category in section.categories:
            f.write(category.wrapped_html_string)
            entry_map = map(wrapString, [i for i in category.entries if (i.date >=start_date) and (i.date <=end_date)])
            entry_str = '\n'.join(entry_map)
            f.write(entry_str)
            #for entry in category.entries:
                #if (entry.date >= start_date) and (entry.date <= end_date):
                    #f.write(entry.wrapped_html_string)
    closing_wrapper = """</body>
    </html>"""
    f.write(closing_wrapper)

    
    
def wrapString(item):
    return item.wrapped_html_string

def getNameValue(item):
    return item.name_value
    
dal = DataAccessLayer()

class App:
    def __init__(self):
        self.d = dal
    
    def setup(self):
        self.d.conn_string = 'sqlite:///roundup_db3.db'
        self.d.connect()
        self.d.session = dal.Session()
    def close(self):
        self.d.session.close()

if __name__ == '__main__':
    pass
#    dal.conn_string = 'sqlite:///roundup_db3.db'
    