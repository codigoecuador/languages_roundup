import itertools as it
import functools
import pprint
import warnings
from collections import Counter
import cmd2 #let's import it so we can pass the cmd interface as an object
import newspaper as np
from dateutil.parser import parse
from sqlalchemy import func
from sqlalchemy.orm.exc import MultipleResultsFound
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import BTCInput3 as btc
from roundup_db1 import Entry, Category, Keyword, Publication, Author, Section, Introduction
from roundup_db1 import DataAccessLayer

class Roundup(object):
    '''NOTE: do not try to use the roundup class for anything except exporting roundups'''
    @classmethod
    def prep_export(cls, title, start_date, end_date,
                   filename, min_category, max_category, intro=None):
        new_sections = get_sections()
        new_sections = [RoundupSection.from_normal_section(i) for i in new_sections]
        result= cls(title=title, start_date=start_date, end_date=end_date,
                   filename=filename, min_category=min_category, max_category=max_category,
                   sections=new_sections, intro=intro)
        result.make_roundup()
        return result
    
    def __init__(self, title, start_date, end_date,
             filename, min_category=1,
             max_category=21, sections=[], intro=None):#app.get_sections()):
        self.title = title
        self.start_date = start_date
        self.end_date = end_date
        self.filename = filename
        self.min_category = min_category
        self.max_category = max_category
        self.sections = [RoundupSection.from_normal_section(i) for i in sections]
        self.categories = self.get_category_articles()
        self.intro=intro
    
    def make_roundup(self):
        '''Takes the sections and gets and organizes the categories. These data structures are used ONLY
        when we export the final roundup to docx or html. There is NO other use for them in the program logic.'''
        for section in self.sections:
            for category in self.categories:
                if category.section_id == section.section_id:
                    section.categories.append(category)
                    #for i in section.categories:
                       # print(len(i.entries))
                    #section.categories = list(it.dropwhile(len(category.entries==0), section.categories))
    
    @staticmethod
    def add_hyperlink(paragraph, text, url):
        # This gets access to the document.xml.rels file and gets a new relation id value
        try:
            part = paragraph.part
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

            # Create the w:hyperlink tag and add needed values
            hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
            hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

            # Create a w:r element and a new w:rPr element
            new_run = docx.oxml.shared.OxmlElement('w:r')
            rPr = docx.oxml.shared.OxmlElement('w:rPr')

            # Join all the xml elements together add add the required text to the w:r element
            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)

        # Create a new Run object and add the hyperlink into it
            r = paragraph.add_run ()
            r._r.append (hyperlink)

        # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
        # Delete this if using a template that has the hyperlink style in it
            r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
            r.font.underline = True

            return hyperlink
        except Exception as e:
            print(e)
            
    @staticmethod
    def add_article(document, article):
    #print(article)
        try:
            new_paragraph = document.add_paragraph('') #add blank paragraph that we append the text to
            Roundup.add_hyperlink(paragraph=new_paragraph, text=article.entry_name, url=article.entry_url)
            #print(Article.get_date_formatted(article))
            #new_paragraph.add_run(' ({0}) '.format(Entry.get_date_formatted(article))) 
            new_paragraph.add_run(f' ({article.get_date_formatted}) ')#.format(article.get_date_formatted))
            #blank space between the link and the description
            new_paragraph.add_run(article.description)
        except Exception as e:
            print(e)
    
    @staticmethod       
    def add_section(document, section):
        document.add_paragraph(section.name)
        section.categories.sort(key=lambda x: x.name, reverse=True)
        section.categories.reverse()
        for category in section.categories:
            #if len(section)
            Roundup.add_category(document, category)
    
    @staticmethod
    def add_category(document, category):
        document.add_paragraph(category.name)
        category.entries.sort(key=lambda x: x.entry_name, reverse=True)
        category.entries.reverse()
        for article in category.entries:
            Roundup.add_article(document, article)
    
    def create_roundup(self, document, roundup_title, sections,
                       intro=None):
        document.add_paragraph(roundup_title)
        if intro != None:
            document.add_paragraph(intro)
        for section in self.sections:
            Roundup.add_section(document, section)
            
    def get_category_articles(self):
        current_category = 1
        max_category = 30
        article_categories = []
        while current_category <= max_category:
    #for category in range(min_category, max_category+1):
            cat = get(session=dal.session, model=Category, category_id=current_category)
        #new_cat = RoundupCategory()
            print(cat)
            try:
                article_categories.append(RoundupCategory(category_id=cat.category_id,
                                                      name=cat.name,
                                                      entries = [i for i in cat.entries if i.date >=self.start_date and i.date <=self.end_date],
                                                         section_id=cat.section_id))
            except AttributeError:
                print('category not found')
            current_category += 1
        return article_categories
        
    def export_docx(self):
        #def complete_roundup2(filename, roundup_title, sections):
        new_document = docx.Document()
        Roundup.create_roundup(self=self, document=new_document,
                               roundup_title=self.title,
                               sections=self.sections,
                               intro =self.intro)
        new_document.save(f'{self.filename}.docx')#.format(self.filename))
        
class RoundupSection(object):
    '''Replicates the section class when we create roundups. DO NOT use this class
    for anything except roundup generation.'''
    @classmethod
    def from_normal_section(cls, section: Section):
        return cls(section_id = section.section_id,
                   name=section.name,
                   categories=[])
    
    def __init__(self, section_id, name, categories=[]):
        self.section_id = section_id
        self.name = name
        self.categories=categories

    def __repr__(self):
        return f'Sec(section_id={self.section_id}, name={self.name})'#.format(self.section_id, self.name)
                   
class RoundupCategory(object):
    '''Replicates the category class when we create roundups. Sqlalchemy methods
    cannot be called on this class. If we did not create a new class type like this,
    then we could not mess with the entries in that category without causing problems
    in the database as a whole. DO NOT use this class for anything except creating
    roundups.'''
    
    @classmethod
    def from_normal_category(cls, category: Category):
        return cls(category_id = category.category_id,
                  name=category.name, section_id = category.section_id, entries = [])
    
    def __init__(self, category_id, name, section_id, entries=[]):
        self.category_id = category_id
        self.name = name
        self.entries = entries
        self.section_id = section_id
        
    def __repr__(self):
        return f'ArtCat(category_id={self.category_id}, name={self.name})'#.format(self.category_id, self.name)

def export_roundup(title, start_date, end_date, filename, intro=None, min_category=1, max_category=21):
    '''Exports a docx roundup'''
    new_roundup = Roundup.prep_export(title=title,
                    start_date=start_date, end_date=end_date,
                    filename=filename,
                    min_category=min_category,
                    max_category=max_category,
                    intro=intro)
    try:
        new_roundup.export_docx()
        print('Export successful')
    except Exception as e:
        print(e)

def create_docx_roundup(args, session=None):
    #we only need a session if we are looking up the introduction
    line=''
    if not args.title:
        title = btc.read_text('Enter title or "." to return to main menu: ')
        if title == '.': return
    else:
        title=' '.join(args.title)
    if not args.introduction:
        if args.introduction_id:
            intro = session.query(Introduction).filter(Introduction.introduction_id==args.introduction_id).first()
            #print(args.introduction_id)
            #print(intro)
            args.introduction = intro.text
            #print(args.introduction)
        else:
            args.introduction=None
    if not args.date_range:
        start_date = btc.read_date('Enter start date "(MM/DD/YYYY)": ')
        end_date = btc.read_date('Enter end_date "MM/DD/YYYY": ')
    else:
        start_date = parse(args.date_range[0]).date()
        end_date = parse(args.date_range[1]).date()
    if not args.filename:
        args.filename = btc.read_text('Enter filename or "." to return to main menu: ')
    #call the export_roundup function to export a docx roundup
    export_roundup(title=title, start_date=start_date,#parse(args.date_range[0]).date(),
                   end_date=end_date, filename=args.filename,
                   intro=args.introduction)
        
#create section

def add_category(session, category_name):
    result = get(session=session, model=Category, name=category_name)
    if result != None:
        print('Category exists')
        return
    else:
        new_category = Category.from_input(category_name=category_name, session=session)
    confirm_choice = btc.read_int_ranged(f'Add {new_category.name}? 1 to add, 2 to cancel: ', 1, 2)
    if confirm_choice == 1:
        session.add(new_category)
        session.commit()
        print(f'{new_category.name} added to database')#.format(new_category.name))
    elif confirm_choice == 2:
        print('Category add cancelled')
        return

def new_cat_with_entry(session, cmdobj, category_name, section_id):
    '''Creates a new category when we're in the process of making an entry'''
    make_cat = lambda x, y: Category(name=x, section_id=y)
    new_category = make_cat(x=category_name, y=section_id)
    try:
        session.add(new_category)
        session.commit()
        #new_cat_style = cmd2.style(text=new_category, bg=cmd2.bg.blue,
        #                           fg=cmd2.fg.white, bold=True)
        #cmdobj.poutput(new_cat_style)
        return None
        #return new_category
    except Exception as e:
        cmdobj.poutput(e)
        warning_msg = cmd2.style(text=f'Category creation failed for {category_name}')
        cmdobj.poutput(warning_msg)
        return None
    

def add_item(session, search_type, new_name):
    """session is the current active session, search_type is the type of item, while
    new_name is the current type of item to add"""
    new_name=new_name.lower()
    search_types = {'author': Author, 'keyword': Keyword,
             'section': Section}
    result = get(session=session, model=search_types[search_type], name=new_name)
    if result != None:
        print('Item exists')
    else:
        confirm_choice = btc.read_int_ranged(f'Add {new_name}? 1 to add, 2 to cancel', 1, 2)
        if confirm_choice == 1:
            new_item=get_or_create(session=session, model=search_types[search_type], name=new_name)
            #session.add(new_keyword)
            #session.commit()
            print(f'{new_item} added to database')#.format(new_item))
        elif confirm_choice == 2:
            print('add cancelled')
            return
    
def add_pub_or_cat(session, add_type, new_name, second_item):
    '''use this one for the CLI'''
    add_type = add_type.lower()
    add_types = {'publication': Publication, 'category': Category}
    #item_twos = {'publication': 'url', 'category': 'section_id'}
    result = get(session=session, model=add_types.get(add_type), name=new_name, second_item=second_item)
    if result != None:
        print(f'{add_type} exists')#.format(add_type))
        print(result)
        return
    else:
        confirm_choice = btc.read_int_ranged(f'Add {new_name}? 1 to add, 2 to cancel', 1, 2)
        if confirm_choice == 1:
            get_or_create(session, model=add_types.get(add_type), name=new_name,
                      second_item=second_item)
        elif confirm_choice == 2:
            print(f'{add_type} add cancelled')#.format(add_type))
            return
        
def add_cat(session, line):
    '''Takes input from the CLI and passes it to the add_pub_or_cat function'''
    try:
        new_name, second_item = line.split()
        add_pub_or_cat(session=session, add_type='category',
                        new_name=new_name, second_item=second_item)
        return True
    except IndexError as e:
        print(e)
        return False

def add_pub(session, line):
    '''Takes input from the CLI and passes it to the add_pub_or_cat function'''
    try:
        new_name, second_item = line.split('|')
        add_pub_or_cat(session=session, add_type='publication',
                        new_name=new_name, second_item=second_item)
        return True
    except IndexError as e:
        print(e)
        return False
        
def add_publication(session, keyword_text):
    result = get(session=session, model=Keyword, word=keyword_text)
    if result != None:
        print('Publication exists')
        return
    else:
        new_keyword = Publication(pub_name=keyword_text)
    confirm_choice = btc.read_int_ranged(f'Add {new_keyword.pub_name}? 1 to add, 2 to cancel', 1, 2)
    if confirm_choice == 1:
        session.add(new_keyword)
        session.commit()
        print(f'{new_keyword.word} added to database')#.format(new_keyword.word))
    elif confirm_choice == 2:
        print('Keyword add cancelled')
        return
    
#modify articles
    
def add_keyword_to_article(session, new_keyword, entry_id=None):##model=Keyword):
    """Add a keyword to an existing article, the keyword is appended to the article's
    keywords attribute"""
    #new_keyword = btc.read_text('Enter new keyword: ')
    #make sure the keyword doesn't already exist
    if entry_id == None:
        entry_id = btc.read_int(prompt='Find an entry by ID: ')
    entry_result = session.query(Entry).filter(Entry.entry_id==entry_id).scalar()
    if entry_result != None:
        print('Entry found: ')
        print(entry_result)
        #new_keyword=btc.read_text('Enter new keyword: ')
        edit_choice = btc.read_int_ranged(f'Add {new_keyword} to this article? (1 for yes, 2 for no): ', 1, 2)
        if edit_choice == 1:
            keyword_result = session.query(Keyword).filter(Keyword.word.like(f'%{new_keyword}%')).all()#.format(new_keyword))).all()
            if len(keyword_result) >= 1:
                print('Keyword exists')
                print(keyword_result)
                print('Entry found:')
                print(entry_result)
                keywords = it.chain(keyword_result)
                while True:
                        #we do this loop if the keyword exists
                    try:
                        item = next(keywords)
                        print(item)
                    except StopIteration:
                        print('No more keywords left')
                    item_choice = btc.read_int_ranged('Is this the keyword you want? (1-yes, 2-continue, 3-quit)', 
                                                      1, 3)
                    #1 select
                    if item_choice == 1:
                        try:
                            assert item not in entry_result.keywords
                        except AssertionError:
                            print('Keyword already attached to article')
                            print('Returning to main menu')
                            return
                        entry_result.keywords.append(item)
                        session.commit()
                        print('Keyword added successfully')
                        break
                    elif item_choice == 2:
                            #continue
                        continue
                    elif item_choice == 3:
                        print('Keyword add cancelled, return to main menu')
                        return
            elif len(keyword_result) ==0:
                print('Keyword does not exist')
                kw = Keyword(word=new_keyword)
                make_keyword_choice = btc.read_int_ranged(f'Create {kw} as a new keyword for ? {entry_result.entry_name} (1 yes, 2 no)',1, 2)
                if make_keyword_choice == 1:
                    entry_result.keywords.append(kw)
                    session.commit()
                    print('Keyword add completed')
                elif make_keyword_choice == 2:
                    print('Add keyword cancelled')
                    return
        elif edit_choice == 2:
            print('Keyword edit cancelled, returning to main menu')
            return
    elif entry_result == None:
        print('Entry not found, returning to main menu')
        return

def delete_entry_keyword(session, entry_id):
    '''Delete a keyword from an existing article by popping it from the keyword list'''
    entry_result = session.query(Entry).filter(Entry.entry_id==entry_id).scalar()
    if entry_result != None:
        article_keywords = it.cycle(entry_result.keywords)
        while True:
            pprint.pprint(entry_result.keywords)
            activeItem = next(article_keywords)
            print('Enter 1 to delete keyword, 2 to continue, 3 to exit to main menu')
            delete_choice = btc.read_int_ranged(f'Delete {activeItem} from the keywords?', 1, 3)
            if delete_choice == 1:
                entry_result.keywords.remove(activeItem)
                session.commit()
            elif delete_choice == 2:
                continue
            elif delete_choice == 3:
                print('Returning to main menu')
                break
    else:
        print('Not found, return to main menu')
        return

    

def make_article(article: str) -> np.Article:
    article = np.Article(article)
    article.build()
    return article

def get_or_create(session, model, **kwargs):
    '''If an object is present in the database, it returns the object.
    Otherwise, it creates a new instance of that object'''
    instance = session.query(model).filter_by(**kwargs).first()
    if instance:
        return instance
    else:
        instance = model(**kwargs)
        session.add(instance)
        session.commit()
        return instance

#read section - functions associated with reading items]

    
def get(session, model, **kwargs):
    '''If an object is present in the database, it returns the object.
    Otherwise, it creates a new instance of that object'''
    instance = session.query(model).filter_by(**kwargs).scalar()
    return instance

def check_date(article):
    '''Check to see if the date is there in the article'''
    try:
        assert article.publish_date, 'No publish date found'
        return article.publish_date
    except AssertionError:
        return None
    
def create_date(article):
    '''Checks to see if the date is there by calling new_date,
    then if there is no date, it asks the user for input'''
    new_date = check_date(article)
    if new_date == None:
        dateObj = input('Enter the date mm/dd/yyyy: ')
        new_date = parse(dateObj)
        #return new_date
    else:
        new_date = input('Enter date (mm/dd/yyyy): ')
        new_date = parse(new_date)
    return new_date    
    
def display_categories(section_id=None):
    query = dal.session.query(Category)
    if section_id != None:
        query = query.filter(Category.section_id == section_id)
    query = query.all()
    print('Categories: ')
    cat_map = map(str, query)
    print('\n'.join(cat_map))

def display_sections(cmdobj=None):
    query = dal.session.query(Section).all()
    section_map = map(str, query)
    if cmdobj==None:
        print('Sections: ')
        print('\n'.join(section_map))
    else:
        section_heading = cmd2.style('Sections:', fg=cmd2.fg.white, bg=cmd2.bg.black, bold=True)
        cmdobj.poutput(section_heading)
        section_str = cmd2.style('\n'.join(section_map), fg=cmd2.fg.white, bg=cmd2.bg.black)
        cmdobj.poutput(section_str)

def get_description():
    description = input('Enter article description (max 500 characters): ')
    return description
        
def get_sections():
    '''Return a list of all the sections in the program'''
    query = dal.session.query(Section).all()
    return query

def get_categories(session, section_id=None):
    '''Return a list of all the categories in the program'''
    query = session.query(Category)
    if section_id == None:
        query = session.query(Category).all()
    else:
        query = query.filter(Category.section_id == section_id).all()
    return query

def count_articles(line, session):
    '''Takes the input from article_count in cmd_roundup.ipynb and determines
    if there is more than one article.'''
    try:
        start_date, end_date = line.split() #tuple unpacking means the first and last elements are taken
        date_range_count(start_date=start_date, end_date=end_date,
                                session=session)
    except ValueError:
        article_count(session=session)

def article_count(session):
    query = session.query(Category.name, func.count(Entry.entry_id))
    query= query.outerjoin(Entry).group_by(Category.name)
    query = query.order_by(func.count(Entry.entry_id))
    query = query.all()
    for row in query[::-1]:
        print(row)
    undescribed_articles=session.query(Entry).filter(Entry.description.like('%not specified%')).all()
    print('Undescribed articles', len(undescribed_articles))
        
def date_range_count(start_date, end_date, session):
    '''Combine date_range_count and article_count'''
    #print(line)
    try:
        #line = line.split(' ')
        start_date, end_date = parse(start_date), parse(end_date)
        start_date = start_date.date()
        end_date = end_date.date()
        print('start date:', start_date)
        print('end date:', end_date)
    except IndexError as e:
            print(e)
            return
    query = session.query(Category.name, func.count(Entry.entry_id))
    query= query.outerjoin(Entry).group_by(Category.name)
    query = query.filter(Entry.date >= start_date, end_date <= end_date)
    #query = query.filter(Entry.date <= end_date)
    query = query.order_by(func.count(Entry.entry_id))
    query = query.all()
    undesc = session.query(Entry).filter(Entry.date >= start_date) #get number of undescribed articles
    undesc = undesc.filter(Entry.date <= end_date)
    undesc = undesc.filter(Entry.description.like('%not specified%')).all()
    undesc_num = len(undesc)
    total = functools.reduce(lambda x, y: x+y,[row[1] for row in query])
    for row in query[::-1]:
        print(row)
    print(total, f'articles total from {start_date} to {end_date}')
    print(f'Undescribed articles: {undesc_num}')
    
def date_range_count2(start_date, end_date, session, cmdobj):
    '''Combine date_range_count and article_count'''
    #print(line)
    try:
        #line = line.split(' ')
        start_date, end_date = parse(start_date), parse(end_date)
        start_date = start_date.date()
        end_date = end_date.date()
        intro=cmd2.style('Entry Count', fg=cmd2.fg.black, bg=cmd2.bg.white, bold=True)
        sdate = cmd2.style(text=f'Start date: {start_date}',fg=cmd2.fg.green, bg=cmd2.bg.white, bold=True)
        edate = cmd2.style(text=f'End date: {end_date}', fg=cmd2.fg.red, bg=cmd2.bg.white, bold=True)
        cmdobj.poutput(intro)
        cmdobj.poutput(sdate)
        cmdobj.poutput(edate)
        #print('start date:', start_date)
        #print('end date:', end_date)
    except IndexError as e:
            print(e)
            return
    ent_query = session.query(Entry)
    ent_query = ent_query.filter(Entry.date >= start_date)
    ent_query = ent_query.filter(Entry.date <= end_date)
    result = ent_query.all()
    mycount = Counter([i.category_name for i in result])    
    pprint.pprint(mycount)
    
def articles_needed(start_date, end_date, session):
    min_articles_cat = 5
    '''Combine date_range_count and article_count'''
    #print(line)
    try:
        #line = line.split(' ')
        start_date, end_date = parse(start_date), parse(end_date)
        start_date = start_date.date()
        end_date = end_date.date()
        print('start date:', start_date)
        print('end date:', end_date)
    except IndexError as e:
            print(e)
            return
    query = session.query(Category.name, func.count(Entry.entry_id))
    query= query.outerjoin(Entry).group_by(Category.name)
    query = query.filter(Entry.date >= start_date, end_date <= end_date)
    #query = query.filter(Entry.date <= end_date)
    query = query.order_by(func.count(Entry.entry_id))
    query = query.all()
    #total = functools.reduce(lambda x, y: x+y,[row[1] for row in query])
    articles_needed = {}
    for row in query[::-1]:
        #print(row)
        if row[1] < 5:
            #print(row, '')
            articles_needed[row[0]] = (min_articles_cat - row[1])
    pprint.pprint(articles_needed)
    #print(articles_needed)
    #for k, v in articles_needed:
        #print('{0}: {1} articles, {2} more articles needed'.format(k, v[0], v[1]))
    #print(total, 'articles total from {0} to {1}'.format(start_date, end_date))
    
def find_introduction(args, session, cmdobj):
    '''Takes arguments from the ui.py module
    args includes args.introduction_id, args.name, and args.text
    only args.introduction_id will be implemented at this present stage'''
    if not args.introduction_id:
        no_intro_id = cmd2.style('No introduction ID found', bg=cmd2.bg.red,
                                 fg=cmd2.fg.white, bold=True)
        raise Exception(no_intro_id)
    else:
        get_intro = lambda x: session.query(Introduction).filter(Introduction.introduction_id==x).one()
        intro_result = get_intro(args.introduction_id)
        result_style = lambda x: cmd2.style(text=x, bg=cmd2.bg.white,
                                            fg=cmd2.fg.black, bold=False)
        cmdobj.poutput(result_style(intro_result))
        
    
def find_entry(args, session, cmdobj):
    if args.entry_id and args.id_range:
        raise Exception('Must be either id or id range')
    elif args.date and args.date_range:
        raise Exception('Must be either date or date range')
    else:
        query = session.query(Entry)
        if args.category_id:
            query = query.filter(Entry.category_id == args.category_id)
        if args.category_name:
            print(args.category_name)
            cat_name = ' '.join(args.category_name)
            cat_query = session.query(Category)
            cat_query = cat_query.filter(Category.name.like(f'%{cat_name}%')).first()
            print(cat_query)
            cat_id = cat_query.id_value
            query = query.filter(Entry.category_id == cat_id)
        if args.id_range:
            query = query.filter(Entry.entry_id >= args.id_range[0],
                                    Entry.entry_id <= args.id_range[1])
        if args.entry_id:
            query = query.filter(Entry.entry_id == args.entry_id)
        if args.date:
            date = parse(args.date).date()
            query = query.filter(Entry.date == date)
        if args.date_range:
            query = query.filter(Entry.date >= parse(args.date_range[0]).date(),
                                    Entry.date <= parse(args.date_range[1]).date())
        if args.url:
            query = query.filter(Entry.url.like(f'%{args.url}%'))
        if args.title:
            print('args.title', args.title)
            query = query.filter(Entry.entry_name.like(f'%{args.title}%'))
        if args.publication_id:
            query = query.filter(Entry.publication_id==args.publication_id)
        if args.publication_title:
            pub_query = session.query(Publication).filter(Publication.title.like(f'%{args.publication_title}%')).first()
            pub_id = pub_query.publication_id
            query=query.filter(Entry.publication_id == pub_id)
            #query=query.filter(Entry.publication.title.like(f'%{args.publication_title}%'))
        result = query.all()
        result_total = len(result)
        if result_total == 0:
            print('no entries found')
            return
        result_cycle = it.cycle(result)
        print(f'{result_total} entries found')
        info_choice = btc.read_int_ranged('1 to view results, 2 to cancel: ', 1, 2)
        if info_choice == 1:
            while True:
                next_item = next(result_cycle)
                continue_choice = btc.read_int_ranged(f'1 to view {next_item.name}, 2 to continue, 3 to quit: ', 1, 3)
                #print(next(result_cycle).name)
                if continue_choice == 1:
                    #print(next(result_cycle))
                    print(next_item)
                    edit_choice = btc.read_int_ranged(f'1-edit {next_item.name}, 2-(continue) 3-quit: ', 1, 3)
                    if edit_choice == 1:
                        edit_entry2(session=session, entry_id=next_item.entry_id, cmdobj=cmdobj)
                    elif edit_choice == 2:
                        continue
                    elif edit_choice == 3:
                        print('Edit cancelled, return to main menu')
                        return
                elif continue_choice == 2:
                    continue
                elif continue_choice == 3:
                    print('returning to main menu')
                    break

def find_section(args, session):
    query = session.query(Section)
    if args.section_id:
        query = query.filter(Section.section_id==args.section_id)
    if args.name:
        query = query.filter(Section.name.like(f'%{args.name}%'))
    result = query.all()
    result_total = len(result)
    if result_total == 0:
        print('no sections found')
        return
    result_cycle = it.cycle(result)
    print(f'{result_total} sections found')
    info_choice = btc.read_int_ranged('1 to view results, 2 to cancel: ', 1, 2)
    if info_choice == 1:
        while True:
            continue_choice = btc.read_int_ranged('1 to view next, 2 to quit: ', 1, 2)
            print(next(result_cycle).name)
            if continue_choice == 1:
                print(next(result_cycle))
            elif continue_choice == 2:
                print('returning to main menu')
                break
                
def find_keyword(args, session):
    query = session.query(Keyword)
    if args.keyword_id:
        query = query.filter(Keyword.keyword_id==args.keyword_id)
    if args.word:
        query = query.filter(Keyword.word.like(f'%{args.word}%'))
    result = query.all()
    result_total = len(result)
    if result_total == 0:
        print('no keywords found')
        return
    result_cycle = it.cycle(result)
    print(f'{result_total} keywords found')
    info_choice = btc.read_int_ranged('1 to view results, 2 to cancel: ', 1, 2)
    if info_choice == 1:
        while True:
            continue_choice = btc.read_int_ranged('1 to view next, 2 to quit: ', 1, 2)
            print(next(result_cycle).name)
            if continue_choice == 1:
                print(next(result_cycle))
            elif continue_choice == 2:
                print('returning to main menu')
                break 

def get_category(category_name, session):
    '''Gets the category for the create article function'''
    category_name = ' '.join(category_name)
    query = session.query(Category).filter(Category.name.like(f'%{category_name}%')).first()
    print(query)
    return query

def get_category_name(category_name, session):
    '''Gets the category for the create article function'''
    #category_name = ' '.join(category_name)
    query = session.query(Category).filter(Category.name.like(f'%{category_name}%')).first()
    #print(query)
    return query

def find_category(args, session):
    query = session.query(Category)
    if args.category_id:
        query = query.filter(Category.category_id == args.category_id)
    if args.category_name:
        print(args.category_name)
        category_name = ' '.join(args.category_name)
        print(category_name)
        query = query.filter(Category.name.like(f"%{category_name}%"))
    if args.section_id:
        query = query.filter(Category.section_id == args.section_id)
    result = query.all()
    result_total = len(result)
    if result_total == 0:
        print('no categories found')
        return
    result_cycle = it.cycle(result)
    print(f'{result_total} categories found')
    info_choice = btc.read_int_ranged('1 to view results, 2 to cancel: ', 1, 2)
    if info_choice == 1:
        while True:
            continue_choice = btc.read_int_ranged('1 to view next, 2 to quit: ', 1, 2)
            print(next(result_cycle).name)
            if continue_choice == 1:
                print(next(result_cycle))
            elif continue_choice == 2:
                print('returning to main menu')
                break

def find_publication(args, session):
    #add publication search
    query = session.query(Publication)
    if args.publication_id:
        query = query.filter(Publication.id_value == args.publication_id)
    if args.title:
        query = query.filter(Publication.title == args.title)
    if args.url:
        query = query.filter(Publication.url.like(f'%{args.url}%'))
    result = query.all()
    print(result)
    results = it.cycle(result)
    active_item = next(results)
    edit_choice = btc.read_int_ranged(prompt=f'Edit publication {active_item.title} (1-y, 2-n)? ',
                                      min_value=1,max_value=2)
    if edit_choice == 1:
        warnings.warn('Edit publication under development')
        edit_single_pub(session=session, active_item=active_item)
    #pass                    

def edit_single_pub(session, active_item):
    '''Edit a single publication'''
    while True:
        print(f'''\n
Publication ID: {active_item.id_value}
Title: {active_item.name_value}
Link: {active_item.url}''')
        edit_choice = btc.read_int_ranged('Edit title - 1, Edit link - 2, 3-next_publication: ', 1, 3)
        if edit_choice == 1:
            view_articles_choice = btc.read_int_ranged('Type 1 to view entries for this publication, 2 to skip', 1, 2)
            if view_articles_choice == 1:
                for i in active_item.entries:
                    print(i)
               # print(f'Entries:\n{active_item.entries}')
            else:
                print('Entries display not needed')
            new_title = btc.read_text('Enter new title or "." to cancel: ')
            print(new_title)
            if new_title != '.': 
                active_item.title = new_title
                session.commit()
            #else:
                #new_desc = 'Not specified' #if the user doens't edit it
        elif edit_choice == 2:
            new_url = btc.read_text('Enter new url or "." to cancel: ')
            if new_url == ".":
                print('Edit cancelled')
                break
            else:
                edit_second_item(session=session, model='publication',
                                 id_value=active_item.id_value,
                                 new_second_value=new_url)
        elif edit_choice == 3:
            break
       
                    
# def search_exact_name(line, session):
#     search_types = {'entry': Entry, 'category': Category, 'publication': Publication,
#                    'section': Section, 'keyword':Keyword}
#     #app.search_exact_name(line=line)
#     line = line.split(' ')
#     search_type = line[0].lower() #make it lowercase to fit in the dictionary
#     value= ' '.join(line[1:]) #The first word is the search type, so we want the remainder of the words
#         #to be the title
#     if search_type in search_types:
#         #print(True)
#         result = get(session=session, model=search_types[search_type], name_value=value)
#         print(result)
#     else:
#         print('Invalid search type, return to main menu')

#def search_by_id(search_type, item_id, session):
#        '''This will serve as a universal function to get an 1 by its id'''
#    search_type = search_type.lower()
#    search_types = {'entry': Entry, 'category': Category,
#                   'publication': Publication, 'section': Section,
#                   'keyword': Keyword, 'author': Author}
#    item_types = {'entry': 'keywords and authors', 'category': 'entries',
#                   'publication': 'entries', 'section': 'categories',
#                   'keyword': 'entries', 'author': 'entries'}
#    if search_type in search_types:
#        result = get(session=session, model=search_types[search_type], id_value=item_id)
#        print(result)
#        info_choice = btc.read_int_ranged('View more information? (1-yes, 2-quit) ',1,2)
#        if info_choice == 1:
#            misc = it.cycle(result.items)
#            while True:
#                print(f'Cycles through {item_types[search_type]}')#.format(item_types[search_type]))
#                continue_choice = btc.read_int_ranged('1 to view next, 2 to quit', 1, 2)
#                if continue_choice == 1:
#                    print(next(misc))
#                elif continue_choice == 2:
#                    print('returning to main menu')
 #                   break
  #  else:
   #     print('Invalid search type. Return to main menu.')

#def name_search(session, line):
#    category_result = session.query(Category).filter(Category.name.like(f'%{line}%'))#.format(line))).all()
#    section_result = session.query(Section).filter(Section.name.like(f'%{line}%'))#.format(line))).all()
#    entry_result = session.query(Entry).filter(Entry.entry_name.like(f'%{line}%'))#.format(line))).all()
#    keyword_result = session.query(Keyword).filter(Keyword.word.like(f'%{line}%')).all()
#    publication_result = session.query(Publication).filter(Publication.title.like(f'%{line}%')).all()
#    author_result = session.query(Author).filter(Author.author_name.like(f'%{line}%')).all()
#    result = it.cycle(category_result+section_result+entry_result+keyword_result+publication_result+author_result)
#    #result2 = it.cycle(result)
#    while True:
#        try:
#            pprint.pprint(next(result))
#        except StopIteration:
#            print('no more left')
#        continue_choice = btc.read_int_ranged('1 to view more results, 2 to return to main menu', 1,2)
#        if continue_choice != 1:
#            break
#            
#def name_search2(session, line):
#    category_result = session.query(Category).filter(Category.name.like(f'%{line}%')).all()
#    section_result = session.query(Section).filter(Section.name.like(f'%{line}%')).all()#.format(line))).all()
#    entry_result = session.query(Entry).filter(Entry.entry_name.like(f'%{line}%')).all()#.format(line))).all()
#    keyword_result = session.query(Keyword).filter(Keyword.word.like(f'%{line}%')).all()#.format(line))).all()
#    publication_result = session.query(Publication).filter(Publication.title.like(f'%{line}%')).all()#like('%{0}%'.format(line))).all()
#    author_result = session.query(Author).filter(Author.author_name.like(f'%{line}%')).all()
#    result = it.cycle(category_result+section_result+entry_result+keyword_result+publication_result+author_result)
#    #result2 = it.cycle(result)
#    while True:
#        try:
#            pprint.pprint(next(result))
#        except StopIteration:
#            print('no more left')
#        continue_choice = btc.read_int_ranged('1 to view more results, 2 to return to main menu', 1,2)
#        if continue_choice != 1:
#            break

def get_articles_for_roundup(start_date, end_date, category_id):
    '''
    Do not mess with this function without absolute certainty that you will
    not break the roundup generation process.
    '''
    query = dal.session.query(Category)
    query = query.filter(Category.category_id == category_id)
    query = query.first()
    return query

def get_entries_by_category(session, line):
    result = session.query(Category)
    result = result.filter(Category.name.like(f'%{line}%'))
    try:
        result = result.one()
    except MultipleResultsFound:
        print('Multiple results found')
        potentials = it.cycle(result)
        while True:
            potential_result = next(potentials)
            print(potential_result)
            result_choice = btc.read_bool(decision='is this the category? (y-yes, n-no): ',
                                         yes='y', no='n',
                                         yes_option='select', no_option='continue')
            if result_choice == True:
                result = potential_result
                break
            else:
                continue
                
    print(result)
    review_choice = btc.read_bool(decision=f'View articles from {result}',
                                       yes='y', no='n',
                                       yes_option='continue', no_option='cancel')
    print(review_choice)
    if review_choice == True:
        entries_by_cat = it.cycle(result.entries)
        while True:
            continue_choice = btc.read_bool(decision=f'View next article from {result}',
                                    yes='y', no='n',
                                    yes_option='continue', no_option='cancel')
            if continue_choice == True:
                print(next(entries_by_cat))
            elif continue_choice == False:
                print('Returning to main menu')
                break

#EDIT Publication
                
def edit_pub(publication_id, session, new_title):
    query=session.query(Publication)
    query=query.filter(Publication.id_value == publication_id)
    result=query.one()
    edit_choice = btc.read_int_ranged(f'Replace title {result.name} with {new_title}: ', 1, 2)
    if edit_choice == 1:
        result.title=new_title
        session.commit()
        return result
    else:
        print('edit cancelled')
        return
    
def edit_pub2(args, session):
    query=session.query(Publication)
    if args.publication_id:
        query=query.filter(Publication.id_value == args.publication_id)
        #result=query.filter(Publication.id_value == publication_id).one()
    elif args.id_range:
        query = query.filter(Publication.id_value >= args.id_range[0],
                             Publication.id_value <= args.id_range[1])
    elif args.title:
        title_value = ' '.join(args.title)
        query=query.filter(Publication.title.like(f"{title_value}"))
    elif args.url:
        query=query.filter(Publication.url.like(f'{args.url}'))
    result = query.all()
    result_total = len(result)
    if result_total == 0:
        print('no publications found')
        return
    result_cycle = it.cycle(result)
    remaining_pubs = len(result)
    print(f'{remaining_pubs} remaining publications')
    while True:
        try:
            active_item=next(result_cycle)
        except StopIteration:
            print('No more publications, return to main menu')
            return
        print('Next publication: ', active_item.title)
        continue_choice = btc.read_int_ranged('press 1 to continue, 2 to quit: ', 1, 2)
        if continue_choice == 1:
            edit_single_pub(session=session, active_item=active_item)
#             #os.system('cls||clear')
#             #undesc = session.query(Entry).filter(Entry.date >= start_date, Entry.date <= end_date)
#             #undesc = undesc.filter(Entry.description.like('%not specified%')).all()
#             #undescribed = len(undesc)
#             #print(f'{undescribed} undescribed articles remaining')
#             #active_item = next(result)
#             while True:
#                 print(f'''\n
# Publication ID: {active_item.id_value}
# Title: {active_item.name_value}
# Link: {active_item.url}''')
#                 edit_choice = btc.read_int_ranged('Edit title - 1, Edit link - 2, 3-next_publication: ', 1, 3)
#                 if edit_choice == 1:
#                     view_articles_choice = btc.read_int_ranged('Type 1 to view entries for this publication, 2 to skip', 1, 2)
#                     if view_articles_choice == 1:
#                         for i in active_item.entries:
#                             print(i)
#                        # print(f'Entries:\n{active_item.entries}')
#                     else:
#                         print('Entries display not needed')
#                     new_title = btc.read_text('Enter new title or "." to cancel: ')
#                     print(new_title)
#                     if new_title != '.': 
#                         active_item.title = new_title
#                         session.commit()
#                     #else:
#                         #new_desc = 'Not specified' #if the user doens't edit it
#                 elif edit_choice == 2:
#                     new_url = btc.read_text('Enter new url or "." to cancel: ')
#                     if new_url == ".":
#                         print('Edit cancelled')
#                         break
#                     else:
#                         edit_second_item(session=session, model='publication',
#                                          id_value=active_item.id_value,
# #                                          new_second_value=new_url)
#                 elif edit_choice == 3:
#                     break
        elif continue_choice == 2:
            break

#Generic EDIT section: these functions relate to editing any item

def edit_name(session, model, id_value, new_name):
    '''Works for: Entry, Category, Section, Keyword, Publication, Author'''
    models = {'entry': Entry, 'category': Category, 'author': Author,
             'section': Section, 'publication': Publication}
    if model in models:
        query = session.query(models.get(model,
            'Please enter "entry", category, section, keyword, publication, author'))
        query = query.filter(models.get(model).id_value == id_value)
        result = query.one()
        confirm_choice = btc.read_int_ranged(f'Replace {result.name_value} with {new_name}? 1-accept, 2-quit', 1, 2)
        if confirm_choice == 1:
            result.name = new_name
            session.commit()
        else:
            print('edit cancelled, return to main menu')
            return
    else:
        return
    
def edit_second_item(session, model, id_value, new_second_value):
    """Works for Publication, Category"""
    models = {'publication': Publication, 'category': Category}
    if model in models:
        query = session.query(models.get(model,
            'Please enter "category" or "publication"'))
        query = query.filter(models.get(model).id_value == id_value)
        result = query.one()
        result.second_item = new_second_value
        session.commit()
    else:
        return
        
        

#Edit category section: these functions relate to editing categories

def edit_category2(args, session):
    '''while loop displaying the current values for the category,
    options for updating, and lets you display the sections. Takes section_id
    and name as optional arguments as well'''
    query=session.query(Category)
    if args.category_id:
        query=query.filter(Category.id_value == args.category_id)
        #result=query.filter(Publication.id_value == publication_id).one()
    elif args.id_range:
        query = query.filter(Category.id_value >= args.id_range[0],
                             Category.id_value <= args.id_range[1])
    elif args.name:
        name_param = ' '.join(args.name)
        query=query.filter(Category.name.like(f"{name_param}"))
    elif args.section_id:
        query=query.filter(Category.section_id)
    result = query.all()
    #KEEP EDITING HERE
    result_total = len(result)
    if result_total == 0:
        print('no publications found')
        return
    result_cycle = it.cycle(result)
    remaining_cats = len(result)
    print(f'{remaining_cats} remaining categories')
    while True:
        try:
            active_item=next(result_cycle)
        except StopIteration:
            print('No more category, return to main menu')
            return
        print('Next category: ', active_item.name)
        continue_choice = btc.read_int_ranged('press 1 to continue, 2 to quit: ', 1, 2)
        if continue_choice == 1:
            #os.system('cls||clear')
            #undesc = session.query(Entry).filter(Entry.date >= start_date, Entry.date <= end_date)
            #undesc = undesc.filter(Entry.description.like('%not specified%')).all()
            #undescribed = len(undesc)
            #print(f'{undescribed} undescribed articles remaining')
            #active_item = next(result)
            while True:
                print(f'''\n
Category ID: {active_item.id_value}
Name: {active_item.name_value}
Section ID: {active_item.section_id}''')
                edit_choice = btc.read_int_ranged('Edit name - 1, Edit section id - 2, 3- next category: ', 1, 3)
                if edit_choice == 1:
                    view_articles_choice = btc.read_int_ranged('Type 1 to view entries for this category, 2 to skip', 1, 2)
                    if view_articles_choice == 1:
                        for i in active_item.entries:
                            print(i)
                       # print(f'Entries:\n{active_item.entries}')
                    else:
                        print('Entries display not needed')
                    new_name = btc.read_text('Enter new name or "." to cancel: ')
                    print(new_name)
                    if new_name != '.': 
                        active_item.name = new_name
                        session.commit()
                    #else:
                        #new_desc = 'Not specified' #if the user doens't edit it
                elif edit_choice == 2:
                    section_list = session.query(Section).all()
                    pprint.pprint(section_list)
                    new_section_id = btc.read_int('Enter new section ID or "0" to cancel: ')
                    if new_section_id == 0:
                        print('Edit cancelled')
                        break
                    else:
                        edit_second_item(session=session, model='category',
                                         id_value=active_item.id_value,
                                         new_second_value=new_section_id)
                elif edit_choice == 3:
                    break
        elif continue_choice == 2:
            break

#Edit keyword section: formatting keywords

#Edit publication section: editing publications

#functions that capture input for a publication

#Edit section section: edit sections

#finalize section: functions to finalize articles
            
def finalize(session, cmdobj, start_date, end_date):
    start_date = parse(start_date)
    end_date = parse(end_date)
    query = session.query(Entry).filter(Entry.date >= start_date, Entry.date <= end_date)
    query = query.filter(Entry.description.like('%not specified%')).all()
    #print(query)
    result = it.cycle(query)
    undescribed_articles = len(query)
    article_count = cmd2.style(text=f'{undescribed_articles} undescribed articles',
                               bg=cmd2.bg.white, fg=cmd2.fg.black)
    cmdobj.poutput(article_count)
    #print(f'{undescribed_articles} undescribed articles')
    while True:
        try:
            active_item=next(result)
        except StopIteration:
            print('No undescribed entries, return to main menu')
            return
        print('Next entry: ', active_item.name)
        continue_choice = btc.read_int_ranged('press 1 to continue, 2 to quit: ', 1, 2)
        if continue_choice == 1:
            #os.system('cls||clear')
            undesc = session.query(Entry).filter(Entry.date >= start_date, Entry.date <= end_date)
            undesc = undesc.filter(Entry.description.like('%not specified%')).all()
            undescribed = len(undesc)
            print(f'{undescribed} undescribed articles remaining')
            #active_item = next(result)
            while True:
                entry_layout_start=cmd2.style(text=f'''\n
Entry ID: {active_item.id_value}
Title: {active_item.name_value}
Date: {active_item.date}''', bg=cmd2.bg.white, fg=cmd2.fg.black, bold=True)

                entry_layout_end=cmd2.style(text=f'''Link: {active_item.entry_url}
Authors: {active_item.authors}
Publication: {active_item.publication}
Category: {active_item.category}
Description: {active_item.description}
Keywords: {active_item.keywords}''', bg=cmd2.bg.white, fg=cmd2.fg.black)
                cmdobj.poutput(entry_layout_start)
                cmdobj.poutput(entry_layout_end)
#                 print(f'''\n
# Entry ID: {active_item.id_value}
# Title: {active_item.name_value}
# Date: {active_item.date}
# Link: {active_item.entry_url}
# Authors: {active_item.authors}
# Publication: {active_item.publication}
# Category: {active_item.category}
# Description: {active_item.description}
# Keywords: {active_item.keywords}''')
                edit_choice = btc.read_int_ranged('Edit description - 1, Edit category id - 2, next_article - 3, quit - 4: ', 1, 4)
                if edit_choice == 1:
                    summary_choice = btc.read_int_ranged('Type 1 to view summary, 2 to skip: ', 1, 2)
                    if summary_choice == 1:
                        summary_heading = cmd2.style(text='Summary:',
                                                     bg=cmd2.bg.white,
                                                     fg=cmd2.fg.black,
                                                     bold=True)
                                                     
                        summary_text=cmd2.style(text=f'{active_item.summary}',
                                                     bg=cmd2.bg.white,
                                                     fg=cmd2.fg.black)
                                                     
                        cmdobj.poutput(summary_heading)
                        cmdobj.poutput(summary_text)
                        #print(f'Summary:\n{active_item.summary}')
                    else:
                        cancelled = cmd2.style(text='Summary display not needed',
                                                     bg=cmd2.bg.white, fg=cmd2.fg.black)
                        cmdobj.poutput(cancelled)
                        #print('Summary display not needed')
                    new_desc = btc.read_text_adv(prompt='Enter new description or "." to cancel: ',
                                                 cmdobj=self, background='blue', foreground='white')
                    cmdobj.poutput(new_desc)
                    #print(new_desc)
                    if new_desc != '.': 
                        active_item.description = new_desc
                        session.commit()
                    else:
                        new_desc = 'Not specified' #if the user doens't edit it
                elif edit_choice == 2:
                    cat_id_finalize(entry_id=active_item.id_value, session=session)
                elif edit_choice == 3:
                    break
                elif edit_choice == 4:
                    raise Exception('Return to main menu')
        elif continue_choice == 2:
            break

#new type of function
#like BTCInput but takes a cmd2obj style as an input
#prints out that style
#then does btc.read_text or read_int or read_int_ranged
#prompt is the prompt for the application

def finalize_two(session, cmdobj, start_date, end_date):
    start_date = parse(start_date)
    end_date = parse(end_date)
    query = session.query(Entry).filter(Entry.date >= start_date, Entry.date <= end_date)
    query = query.filter(Entry.description.like('%not specified%')).all()
    #print(query)
    result = it.cycle(query)
    undescribed_articles = len(query)
    article_count = cmd2.style(text=f'{undescribed_articles} undescribed articles',
                               bg=cmd2.bg.white, fg=cmd2.fg.black)
    cmdobj.poutput(article_count)
    #print(f'{undescribed_articles} undescribed articles')
    while True:
        try:
            active_item=next(result)
        except StopIteration:
            print('No undescribed entries, return to main menu')
            return
        #get the number of undescribed entries
        #print the number of undescribed entries with proper style
        #print article layout
        #take user choice
        undesc = session.query(Entry).filter(Entry.date >= start_date, Entry.date <= end_date)
        undesc = undesc.filter(Entry.description.like('%not specified%')).all()
        undescribed = len(undesc)
        print(f'{undescribed} undescribed articles remaining')
        #active_item = next(result)
        while True:
            entry_layout_start=cmd2.style(text=f'''\n
Entry ID: {active_item.id_value}
Title: {active_item.name_value}
Date: {active_item.date}''', bg=cmd2.bg.white, fg=cmd2.fg.black, bold=True)

            entry_layout_end=cmd2.style(text=f'''Link: {active_item.entry_url}
Authors: {active_item.authors}
Publication: {active_item.publication}
Category: {active_item.category}
Description: {active_item.description}
Keywords: {active_item.keywords}''', bg=cmd2.bg.white, fg=cmd2.fg.black)
            cmdobj.poutput(entry_layout_start)
            cmdobj.poutput(entry_layout_end)
            edit_choice = btc.read_int_ranged_adv(prompt='Edit description - 1, Edit category id - 2, next_article - 3, quit - 4: ',
                                                  min_value=1, max_value=4,
                                                  cmdobj=cmdobj,
                                                  fg=cmd2.fg.white,
                                                  bg=cmd2.bg.blue,
                                                  bold=True)
            if edit_choice == 1:
                summary_heading = cmd2.style(text='Summary:',
                                             bg=cmd2.bg.white,
                                             fg=cmd2.fg.black,
                                             bold=True)
                                             
                summary_text=cmd2.style(text=f'{active_item.summary}',
                                             bg=cmd2.bg.white,
                                             fg=cmd2.fg.black)
                                             
                cmdobj.poutput(summary_heading)
                cmdobj.poutput(summary_text)
                new_desc = btc.read_text_adv(prompt='Enter new description or "." to cancel, "`" to erase: ',
                                             cmdobj=cmdobj, bold=False)# bg=cmd2.bg.blue, fg=cmd2.fg.white)
                #cmdobj.poutput(new_desc)
                #print(new_desc)
                if new_desc == '.':
                    cancel_message =  cmd2.style(text='edit cancelled',
                                                    bg=cmd2.bg.blue, fg=cmd2.fg.white, bold=True)
                    cmdobj.poutput(cancel_message)
                    #break
                elif new_desc == '`':
                    erase_message =  cmd2.style(text='edit cancelled',
                                                    bg=cmd2.bg.blue, fg=cmd2.fg.white, bold=True)
                    cmdobj.poutput(erase_message)
                    #new_desc = 'Not specified'
                    active_item.description = 'Not specified'
                    session.commit()
                else:
                    #warnings.warn('Cut and paste the description again or it will be erased')
                    active_item.description = new_desc
                    session.commit()
                #else:
                #    new_desc = 'Not specified' #if the user doens't edit it
            elif edit_choice == 2:
                cat_id_finalize(entry_id=active_item.id_value, session=session)
            elif edit_choice == 3:
                break
            elif edit_choice == 4:
                #raise Exception('Return to main menu')
                return_message = cmd2.style(text='Return to main menu',
                                                bg=cmd2.bg.red,
                                                fg=cmd2.fg.yellow, bold=True)
                #cmdobj.poutput(return_message)
                raise Exception(return_message)
                    


def set_pubnames(args, session):
    if args.id_range:
        minimum_id = args.id_range[0]
        maximum_id = args.id_range[1]
        result = session.query(Publication).filter(Publication.id_value >= minimum_id, 
                                  Publication.id_value <= maximum_id).all()
        #result = result.filter(Publication.url == Publication.title).all()
        results = it.cycle([i for i in result if i.url == i.name_value])
        while True:
            
            new_pub = next(results)
            edit_choice = btc.read_int_ranged(f'Type 1 to edit {new_pub.title}, 2 to continue, 3 to quit: ', 1, 3)
            if edit_choice == 1:
                new_title = btc.read_text('Enter new title or "." to cancel: ')
                if new_title != ".":
                    new_pub.title = new_title
                    session.commit()
                else:
                    print('Edit cancelled')
                    #break
            elif edit_choice == 3:
                print('Edit cancelled')
                break
                #break
                
def merge_pub(args, session):
    #args contains: url
    #get all publication urls:
    query_one = session.query(Publication).all()
    if args.id_range:
        query_one = query_one.filter(Publication.id_value >= args.id_range[0],
                                     Publication.id_value <= args.id_range[1])
    duplicate_counter = Counter([i.url for i in query_one])
    for i in duplicate_counter.items():
        print(i)
    
 #   if args.url:
     #   query = session.query(Publication.url.like(f'%{args.url}%')).all()
    #    results = it.cycle(query)
     #   pass
    #else:
      #  return


#Edit entry section: these functions relate to editing entries

def name_from_input(session, entry_id):
    '''Gets user input for the article name, the name editing itself is carried out by the edit_name function'''
    query = session.query(Entry)
    query = query.filter(Entry.entry_id == entry_id)
    result = query.one()
    print(result)
    edit_choice = btc.read_int_ranged('Edit title (1 for yes, 2 for no): ', 1, 2)
    if edit_choice == 1:
        new_name = btc.read_text('Enter new title or "." to cancel: ')
        if new_name != '.':
            #edit_name(session=session, entry_id=entry_id, new_name=new_name)
            edit_name(session=session, model='entry', id_value = entry_id, new_name=new_name)
        else:
            print('Edit description cancelled')
            return
    if edit_choice == 2:
        print('Edit cancelled')
        return
    
def edit_date(session, entry_id, new_date):
    query = session.query(Entry)
    query = query.filter(Entry.entry_id == entry_id)
    result = query.one()
    result.date = new_date
    session.commit()
    print('Entry edit successful')
    
def date_from_input(session, entry_id):
    query = session.query(Entry)
    query = query.filter(Entry.entry_id == entry_id)
    result = query.one()
    print(result)
    edit_choice = btc.read_int_ranged('Edit date (1 for yes, 2 for no): ', 1, 2)
    if edit_choice == 1:
        new_date = btc.read_text('Enter new date or "." to cancel: ')
        new_date = parse(new_date)
        if new_date != '.':
            edit_date(session=session, entry_id=entry_id, new_date=new_date)
        else:
            print('Edit date cancelled')
            return
    if edit_choice == 2:
        print('Edit cancelled')
        return
    
def edit_description(session, entry_id, new_description):
    query = session.query(Entry)
    query = query.filter(Entry.entry_id == entry_id)
    result = query.one()
    result.description = new_description
    session.commit()
    print('Entry edit successful')
    
def desc_from_input(session, entry_id):
    query = session.query(Entry)
    query = query.filter(Entry.entry_id == entry_id)
    result = query.one()
    print(result)
    edit_choice = btc.read_int_ranged('Edit description (1 for yes, 2 for no): ', 1, 2)
    if edit_choice == 1:
        new_description = btc.read_text('Enter new description or "." to cancel: ')
        if new_description != '.':
            edit_description(session=session, entry_id=entry_id, new_description=new_description)
        else:
            print('Edit description cancelled')
            return
    if edit_choice == 2:
        print('Edit cancelled')
        return
    
def desc_finalize(session, entry_id):
    edit_choice = btc.read_int_ranged('Edit description (1 for yes, 2 for no): ', 1, 2)
    if edit_choice == 1:
        new_description = btc.read_text('Enter new description or "." to cancel: ')
        if new_description != '.':
            edit_description(session=session, entry_id=entry_id, new_description=new_description)
        else:
            print('Edit description cancelled')
            return
    if edit_choice == 2:
        print('Edit cancelled')
        return

def cat_id_finalize(session, entry_id):
    display_categories()
    new_category_id = btc.read_int('Enter new category id or 0 to cancel: ')
    if new_category_id == 0:
        print('edit category id cancelled')
    else:
        query = session.query(Entry)
        query = query.filter(Entry.entry_id == entry_id)
        result = query.one()
        result.category_id = new_category_id
        session.commit()
        print('Entry edit successful')    
    
def edit_category_id(session, entry_id, new_category_id):
    query = session.query(Entry)
    query = query.filter(Entry.entry_id == entry_id)
    result = query.one()
    result.category_id = new_category_id
    session.commit()
    print('Entry edit successful')
    
def cat_id_from_input(session, entry_id):
    query = session.query(Entry)
    query = query.filter(Entry.entry_id==entry_id)
    result = query.one()
    print(result)
    edit_choice=btc.read_int_ranged('Edit category ID (1 for yes, 2 for no): ', 1, 2)
    if edit_choice == 1:
        display_categories()
        new_category_id = btc.read_int('Enter category id or 0 to cancel: ')
        if new_category_id <= 0:
            print('Edit category ID cancelled, return to main menu')
            return
        else:
            edit_category_id(session=session, entry_id=entry_id,
                             new_category_id=new_category_id)
            
        
def edit_entry(session, entry_id, cmdobj):
    entry = get(session=session, model=Entry, entry_id=entry_id)
    choices = {1: name_from_input, 2: date_from_input,
               3: cat_id_from_input, 4: desc_from_input}
    #print(entry)
    while True:
        e_style = cmd2.style(entry, fg=cmd2.fg.white, bg=cmd2.bg.blue, bold=False)
        cmdobj.poutput(e_style)
        #print(entry)
        choice = btc.read_int_ranged_adv(f'''Continue editing {entry.entry_name}?
1 - edit name
2 - edit date
3 - edit category id
4 - edit description
5 - finish editing
Enter selection: ''', cmdobj=cmdobj, min_value=1, max_value=5, bold=True)
        if choice <= 4:
            choices.get(choice)(session, entry_id)
        elif choice == 5:
            editing_complete = cmd2.style('Editing complete', bg=cmd2.bg.red,
                                          fg=cmd2.fg.white, bold=True)
            cmdobj.poutput(editing_complete)
            #print('Editing complete')
            break
    #pass
#create section - functions for creating new instances

#add_articles: this section is for adding entries, this is separate due to the complexity of the code. The entries
#must be downloaded using the newspaper app.

#def add_entry(session, url, category_id=None, date=None):
#    new_entry = from_newspaper_two(url=url, category_id=category_id, date=date)
#    session.add(new_entry)
#    session.commit()
#    print(f'{new_entry.entry_name} added to database')#.format(new_entry.entry_name))

def add_intro(session, cmdobj, name, text):
    '''Add introduction to database from a user command'''
    new_intro = Introduction(name=name, text=text)
    session.add(new_intro)
    session.commit()
    add_complete_intro = cmd2.style(text=f'{new_intro}')
    cmdobj.poutput(add_complete_intro)
    

def add_entry(session, cmdobj, url, category_id, date, description=None, category_name=None,
              manual_description=False, new_keywords=None):
    'qa is short for quick_add'
    try:
        #assert category_id.isnumeric(), 'category_id out of order'
        #assert parse(date), 'date out of order'
        assert len(url) > 10, 'url out of order'
    
    except AssertionError as e:
        #one of the parameters is entered incorrectly
        print(e)
        return
    new_article = make_article(url)
    add_notice = cmd2.style('\n Title is being added...', fg=cmd2.fg.white,
                            bg=cmd2.bg.blue)
    cmdobj.poutput(add_notice)
    #print('\nTitle is being added...')#, new_article.title)
    article_title = cmd2.style(text=new_article.title, fg=cmd2.fg.white,
                               bg=cmd2.bg.blue, bold=True)
    cmdobj.poutput(article_title)
    #pprint.pprint(new_article.title)
    new_pub = get(session, Publication,
                               url=new_article.source_url)
    new_pub_notice = cmd2.style(text=new_pub, fg=cmd2.fg.yellow,
                                bg=cmd2.bg.green, bold=True)
    cmdobj.poutput(new_pub_notice)
    #print(new_pub)
    if new_pub == None:
        new_pub = get_or_create(session, Publication,
                        title=new_article.source_url,
                        url=new_article.source_url)
        cmdobj.poutput(new_pub_notice)
        #print(new_pub)
    if type(date) == str:
        date=parse(date)
    if description == None:
        if manual_description == True:
            pprint.pprint(new_article.title)
            pprint.pprint(new_article.summary)
            description_prompt = btc.read_text('Enter article description or "." to cancel: ')
            if description_prompt == '.':
                description='Not specified'
        elif manual_description==False:
            description='Not specified'
    if new_keywords:
        for i in new_keywords:
            #if type(i) == list:
                #raise Exception('manually added keywords must be a single word')
                #i = ' '.join(i)
            choice = btc.read_int_ranged(f'add {i} to keywords?', 1, 2)
            if choice == 1:
                new_article.keywords.append(i)
                print(f'{i} added to keywords manually')
    authors = [get_or_create(session, Author, author_name=i) for i in new_article.authors]
    keywords = [get_or_create(session, Keyword, word=i) for i in new_article.keywords]
    
    new_entry= create_entry(article=new_article, description=description,
                            publication_id=new_pub.publication_id, category_id=category_id,
                            date=date, authors=authors, keywords=keywords)
    session.add(new_entry)
    session.commit()
    #print(f'{new_entry}')
    new_entry_format = cmd2.style(text=str(new_entry), bg=cmd2.bg.blue,
                                  fg=cmd2.fg.white, bold=False)
    cmdobj.poutput(new_entry_format)
    success = cmd2.style(text='added successfully', bg=cmd2.bg.blue,
                         fg=cmd2.fg.white, bold=True)
    cmdobj.poutput(success)
    #print(f'{new_entry.name}, entry id {new_entry.id_value} added successfully')
    
#def from_newspaper(url):
#    #include confirm option to make sure that the user wants to add the article
#    new_article = make_article(url)
#    print('\nTitle:')#, new_article.title)
#    pprint.pprint(new_article.title)
#    print('Summary:')#, new_article.summary)
#    pprint.pprint(new_article.summary)
#    try:
#        new_pub = get(dal.session, Publication,
#                               url=new_article.source_url)
#        if new_pub != None:
#            print(new_pub)
#            pub_choice = btc.read_bool(decision=f'Confirm {new_pub} as article publication? ',
#                                       yes='y', no='n',
#                                       yes_option='confirm', no_option='cancel')
#            if pub_choice != True:
#                #if the user REJECTS the publication that is listed
#                new_title = input('Enter publication title: ')
#                new_source_url = input('Enter source URL: ')
#                new_pub = get_or_create(dal.session, Publication,
#                                title=new_title,
#                               url=new_source_url)
#        else:
#            #if we have a source URL but no title
#            new_title = input('Enter publication title: ')
#            new_pub = get_or_create(dal.session, Publication,
#                                title=new_title,
#                               url=new_article.source_url)
#    except Exception as e:
#        print(e)
#        #if there's already a publication with that URL
#        new_title = input('Enter publication title: ')
#        new_source_url = input('Enter source URL: ')
#        new_pub = get_or_create(dal.session, Publication,
#                                title=new_title,
#                               url=new_source_url)
#    publication_id = new_pub.publication_id
#    print('-'*64)
#    display_categories()
#    category_id = int(input('Enter category ID: '))
#    #pub_title = input('Enter publication title: ')
#    
#    date = create_date(new_article)
#    description=get_description()
#    confirm_choice = btc.read_int_ranged('Confirm article add (1-yes, 2-no) : ', 1, 2)
#    if confirm_choice == 2:
#        print('Article add cancelled')
#        return
#    else:
#        authors = [get_or_create(dal.session, Author, author_name=i) for i in new_article.authors]
#        keywords = [get_or_create(dal.session, Keyword, word=i) for i in new_article.keywords]
#        return create_entry(article=new_article, description=description,
#                            publication_id=publication_id, category_id=category_id,
#                            date=date, authors=authors, keywords=keywords)
#    
#def from_newspaper_two(url, category_id = None, date=None):
#    #include confirm option to make sure that the user wants to add the article
#    new_article = make_article(url)
#    print('\nTitle:')#, new_article.title)
#    pprint.pprint(new_article.title)
#    print('Summary:')#, new_article.summary)
#    pprint.pprint(new_article.summary)
#    try:
#        new_pub = get(dal.session, Publication,
#                               url=new_article.source_url)
#        if new_pub != None:
#            print(new_pub)
#            pub_choice = btc.read_bool(decision=f'Confirm {new_pub} as article publication? ',
#                                       yes='y', no='n',
#                                       yes_option='confirm', no_option='cancel')
#            if pub_choice != True:
#                #if the user REJECTS the publication that is listed
#                new_title = input('Enter publication title: ')
#                new_source_url = input('Enter source URL: ')
#                new_pub = get_or_create(dal.session, Publication,
#                                title=new_title,
#                               url=new_source_url)
#        else:
#            #if we have a source URL but no title
#            new_title = input('Enter publication title: ')
#            new_pub = get_or_create(dal.session, Publication,
#                                title=new_title,
#                               url=new_article.source_url)
#    except Exception as e:
#        print(e)
#        #if there's already a publication with that URL
#        new_title = input('Enter publication title: ')
#        new_source_url = input('Enter source URL: ')
#        new_pub = get_or_create(dal.session, Publication,
#                                title=new_title,
#                               url=new_source_url)
#    publication_id = new_pub.publication_id
#    print('-'*64)
#    if category_id == None:
#        display_categories()
#        category_id = int(input('Enter category ID: '))
#    #pub_title = input('Enter publication title: ')
#    if date == None:
#        date = create_date(new_article)
#    description=get_description()
#    confirm_choice = btc.read_int_ranged('Confirm article add (1-yes, 2-no) : ', 1, 2)
#    if confirm_choice == 2:
#        print('Article add cancelled')
#        return
#    else:
#        authors = [get_or_create(dal.session, Author, author_name=i) for i in new_article.authors]
#        keywords = [get_or_create(dal.session, Keyword, word=i) for i in new_article.keywords]
#        return create_entry(article=new_article, description=description,
#                            publication_id=publication_id, category_id=category_id,
#                            date=date, authors=authors, keywords=keywords)



def create_entry(article, description, publication_id, category_id, date, authors=[],
                 keywords=[]):
    #date=parse(date)
    return Entry(entry_name=article.title, entry_url=article.url,
          description=description, summary=article.summary,
        authors=authors,
          publication_id=publication_id,
    category_id=category_id,
        keywords=keywords,
            date=date)

#delete section - delete article

def delete_item(session, cmdobj, model, id_value):
    model = model.lower()
    models = {'entry': Entry, 'category': Category, 'keyword': Keyword,
             'author': Author, 'publication': Publication, 'section': Section}
    result = get(session=session, model = models.get(model, 'invalid delete type'),
                id_value=id_value)
    if result != None:
        result_style = cmd2.style(text=result, fg=cmd2.fg.white, bg=cmd2.bg.blue,
                                  bold=False)
        cmdobj.poutput(result_style)
        delete_choice = btc.read_int_ranged_adv(f'Delete {model} (1 for yes, 2 for no)?',
                                                min_value=1, max_value=2,
                                                fg=cmd2.fg.white,
                                                bg=cmd2.bg.red,
                                                bold=True,
                                                cmdobj=cmdobj)         
        if delete_choice == 1:
            if (model == 'category') or (model=='section') or (model=='publication'):
                try:
                    assert len(result.items) == 0
                except AssertionError:
                    has_items_message = cmd2.style('Result has items, delete these first ',
                                                   fg=cmd2.fg.yellow, bg=cmd2.bg.red,
                                                   bold=True,cmdobj=cmdobj)
                    cmdobj.poutput(has_items_message)
                    cmdobj.poutput(result.items)
                    #print('result has items, delete these first: ', result.items)
                    return
            confirm_choice = btc.read_int_ranged_adv('Are you sure (1 for yes, 2 for no)?',
                                                     min_value=1, max_value=2,
                                                     fg=cmd2.fg.yellow,
                                                     bg=cmd2.bg.red,
                                                     bold=True,cmdobj=cmdobj)
            if confirm_choice == 1:
                #delete the article
                session.delete(result)
                session.commit()
                delete_message = cmd2.style(f'{result_style} deleted',
                                            fg=cmd2.fg.yellow,
                                            bg=cmd2.bg.red,
                                            bold=True)
                cmdobj.poutput(delete_message)
                #print(f'{model} deleted')#.format(model))
            elif confirm_choice == 2:
                cancel_message= cmd2.style('Delete cancelled',
                                           fg=cmd2.fg.yellow,
                                           bg=cmd2.bg.red,
                                           bold=True)
                remains_message = cmd2.style('f{result.name_value} remains in database',
                                             fg=cmd2.fg.yellow,
                                             bg=cmd2.bg.red,
                                             bold=False)
                cmdobj.poutput(cancel_message)
                cmdobj.poutput(remains_message)
                #print('Delete cancelled')
                #print(f'{result.name_value} remains in database')#.format(result.name_value))
        else:
            delete_cancelled = cmd2.style('Delete cancelled by user, returning to main menu',
                                          bg=cmd2.bg.red, fg=cmd2.fg.yellow,
                                          bold=True)
            cmdobj.poutput(delete_cancelled)
            #print('Delete cancelled by user, returning to main menu')
    else:
        not_found = cmd2.style('Item not found, delete cancelled',
                               bg=cmd2.bg.red, fg=cmd2.fg.white,
                               bold=False)
        cmdobj.poutput(not_found)
        #print('Item not found, delete cancelled')

def make_html_roundup(line, session):
    del line
    filename=btc.read_text('Please enter filename or "." to cancel: ')
    if filename == ".": return
    title=btc.read_text('Please enter title or "." to cancel: ')
    if title == ".": return
    start_date = btc.read_date('Pease enter start date ("MM/DD/YYYY"): ')
    end_date = btc.read_date('Please enter end date ("MM/DD/YYYY"):')
    try:
        export_html2(session=session, program=filename, title=title,
                           start_date=start_date, end_date=end_date)
        print(f'{filename} exported successfully')#.format(filename) )
    except Exception as e:
        print(e)
    
def export_html2(session, program, start_date, end_date, title):
    filename = program + '.html'
    f = open(filename, 'w')

    opening_wrapper = f"""<html>
    <head>
    <title>{title}</title>
    </head>
    <body><p>{title}</p>"""#.format(title)
    f.write(opening_wrapper)
    section_query = session.query(Section)
    section_query = section_query.all()
    for section in section_query:
        f.write(section.wrapped_html_string)
        for category in section.categories:
            f.write(category.wrapped_html_string)
            entry_map = map(wrapString, [i for i in category.entries if (i.date >=start_date) and (i.date <=end_date)])
            entry_str = '\n'.join(entry_map)
            f.write(entry_str)
    closing_wrapper = """</body>
    </html>"""
    f.write(closing_wrapper)

def export_jsx(session, program, start_date, end_date, title):
    """Do not mess with this function it exports roundups"""
    filename = program + '.html'
    f = open(filename, 'w')

    opening_wrapper = f"""<html>
    <head>
    <title>{title}</title>
    </head>
    <body><p>{title}</p>"""#.format(title)
    f.write(opening_wrapper)
    section_query = session.query(Section)
    section_query = section_query.all()
    for section in section_query:
        f.write(section.wrapped_jsx_string)
        for category in section.categories:
            f.write(category.wrapped_jsx_string)
            entry_map = map(wrapStringJSX, [i for i in category.entries if (i.date >=start_date) and (i.date <=end_date)])
            entry_str = '\n'.join(entry_map)
            f.write(entry_str)
            #for entry in category.entries:
                #if (entry.date >= start_date) and (entry.date <= end_date):
                    #f.write(entry.wrapped_html_string)
    closing_wrapper = """</body>
    </html>"""
    f.write(closing_wrapper)
    
def exp_full_jsx(session, program, start_date, end_date, title, use_sections=False):
    """Do not mess with this function, it exports roundups"""
    opening_wrapper = """import React, { Component } from "react";
import { connect } from "react-redux";
import "slick-carousel/slick/slick.css";
import "slick-carousel/slick/slick-theme.css";
import "./Roundup.css";

class RoundupContainer extends Component {
  render() {
    let num;
    this.props.size === "mobile" ? (num = 1) : (num = 2);

    return (
      <div className="main-container">
        <div className="headline roundup-headline">
          <span className="gold">Roundup</span>
        </div>

        <div className="roundup-container-text">
          <div className="subhead center roundup-subhead">
            News Roundup: latest updates on coding and web development
          </div>

          <p>"""
          
    closing_wrapper = '''</p>
        </div>

      </div>
    );
  }
}

const mapStateToProps = state => {
  return { size: state.size };
};
export default connect(mapStateToProps)(RoundupContainer);'''
    
    #export a full JSX file
    #if use_sections == True:
       # warnings.warn('Sections for JSX not implemented yet')
       # f.close()
       # return
    filename = program + '.js'
    
    with open(filename, 'w') as f:
        f.write(opening_wrapper)
        if use_sections == False:
            cat_query = session.query(Category).all()
            for cat in cat_query: #new way of displaying the category names
                f.write(f"<p><b>{cat.name.title()}</b></p>\n")
                entry_map = map(wrapStringJSX, [i for i in cat.entries if (i.date >=start_date) and (i.date <=end_date)])
                entry_str = '\n'.join(entry_map)
                f.write(entry_str)
        elif use_sections == True:
            print('sections included in roundup')
            section_query = session.query(Section)
            section_query = section_query.all()
            for section in section_query:
                f.write(section.wrapped_jsx_string)
                for category in section.categories:
                    f.write(category.wrapped_jsx_string)
                    entry_map = map(wrapStringJSX, [i for i in category.entries if (i.date >=start_date) and (i.date <=end_date)])
                    entry_str = '\n'.join(entry_map)
                    f.write(entry_str)
        f.write(closing_wrapper)
    
def wrapString(item):
    return item.wrapped_html_string

def wrapStringJSX(item):
    return item.wrapped_jsx_string

def getNameValue(item):
    return item.name_value
    
dal = DataAccessLayer()

class App:
    def __init__(self):
        self.d = dal
    
    def setup(self):
        self.d.conn_string = 'sqlite:///roundup_db3.db'
        self.d.connect()
        self.d.session = dal.Session()
    def close(self):
        self.d.session.close()

if __name__ == '__main__':
    pass
#    dal.conn_string = 'sqlite:///roundup_db3.db'
    